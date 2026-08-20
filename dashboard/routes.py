"""
dashboard/routes.py

Tableau de bord personnalisé : service, direction et admin (éditeur + lecteur).
Statistiques PTA sans montants : comptages, taux d'exécution, statuts.
Admin → vue globale avec filtres direction / service + ventilation par nature.
"""
import io
import os
import datetime
from flask import render_template, request, session, redirect, url_for, flash, send_file, current_app
from flask_login import login_required, current_user
from models import Annee, Service, Direction, Programme
from dashboard import dashboard_bp

INVEST = "Activité d'investissement"


# ─────────────────────────────────────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _get_annee():
    annee_id = session.get('annee_id')
    if annee_id:
        return Annee.query.get(annee_id)
    return Annee.query.filter_by(actif=True).first()


def _cibles_service(annee, service):
    """Cibles théoriques T1-T4 pour un service, identiques à celles du module Stats."""
    from stats.routes import _compute_cibles_services
    result = _compute_cibles_services(annee, [service])
    if result:
        return result[0]['cibles']
    return {1: 0.0, 2: 0.0, 3: 0.0, 4: 0.0}


def _cibles_direction(annee, direction):
    """Cibles théoriques T1-T4 pour une direction, identiques à celles du module Stats."""
    from stats.routes import _compute_cibles_directions
    result = _compute_cibles_directions(annee, [direction])
    if result:
        return result[0]['cibles']
    return {1: 0.0, 2: 0.0, 3: 0.0, 4: 0.0}


def _compute_synthese_admin(annee):
    """
    Pour les admins : calcule le taux global de chaque direction et de chacun
    de ses services. Retourne une liste triée par direction.
    Partage le même suivi_map (un seul appel DB) pour toutes les entités.
    """
    from dirpta.routes import _compute_pta_direction
    from svcpta.routes import _compute_pta_service
    from suivi.routes import _filter_and_renorm, _enrich, _load_suivis_global

    suivi_map  = _load_suivis_global(annee.id)
    directions = Direction.query.order_by(Direction.nom).all()
    synthese   = []

    for direction in directions:
        # ── Taux de la direction (vue globale, trim=0) ──────────────────────
        data_dir = _compute_pta_direction(annee, direction)
        data_g   = _filter_and_renorm(data_dir, 0)
        if data_g:
            taux_dir, _ = _enrich(data_g, suivi_map, None)
            nb_t_dir    = sum(1 for pd in data_g
                              for pjd in pd['projets']
                              for ad  in pjd['activites']
                              for td  in ad['taches'])
        else:
            taux_dir, nb_t_dir = 0.0, 0

        # ── Services de cette direction ─────────────────────────────────────
        services  = Service.query.filter_by(direction_id=direction.id)\
                                 .order_by(Service.nom).all()
        svcs_list = []
        for svc in services:
            data_svc = _compute_pta_service(annee, svc)
            data_sg  = _filter_and_renorm(data_svc, 0)
            if data_sg:
                taux_svc, _ = _enrich(data_sg, suivi_map, None)
                nb_t_svc    = sum(1 for pd in data_sg
                                  for pjd in pd['projets']
                                  for ad  in pjd['activites']
                                  for td  in ad['taches'])
            else:
                taux_svc, nb_t_svc = 0.0, 0
            svcs_list.append({
                'service':   svc,
                'taux':      taux_svc,
                'nb_taches': nb_t_svc,
            })

        synthese.append({
            'direction': direction,
            'taux':      taux_dir,
            'nb_taches': nb_t_dir,
            'services':  svcs_list,
        })

    return synthese


def _cibles_global(annee):
    """Cibles théoriques T1-T4 pour le PTA global (tous programmes)."""
    from stats.routes import _compute_cibles
    programmes = Programme.query.filter_by(annee_id=annee.id).all()
    if not programmes:
        return {1: 0.0, 2: 0.0, 3: 0.0, 4: 0.0}
    result = _compute_cibles(programmes)
    return result.get('global', {1: 0.0, 2: 0.0, 3: 0.0, 4: 0.0})


def _compute_dashboard_stats(annee, data, cibles, with_nature=False):
    """
    Calcule toutes les statistiques du tableau de bord.
    cibles      : dict {1: %, 2: %, 3: %, 4: %} issu du module Stats.
    with_nature : si True, calcule aussi suivi_invest et suivi_fonct (admins).
    """
    from suivi.routes import _filter_and_renorm, _enrich, _load_suivis_global

    # ── Comptages PTA ────────────────────────────────────────────────────────
    all_acts   = [ad for pd in data
                  for pjd in pd['projets'] for ad in pjd['activites']]
    all_taches = [td for ad in all_acts for td in ad['taches']]

    nb_programmes = len(data)
    nb_projets    = sum(len(pd['projets']) for pd in data)
    nb_activites  = len(all_acts)
    nb_taches     = len(all_taches)

    acts_invest = [ad for ad in all_acts
                   if (ad['activite'].type_activite or '') == INVEST]
    acts_fonct  = [ad for ad in all_acts
                   if (ad['activite'].type_activite or '') != INVEST]

    nb_invest        = len(acts_invest)
    nb_fonct         = len(acts_fonct)
    nb_taches_invest = sum(len(ad['taches']) for ad in acts_invest)
    nb_taches_fonct  = sum(len(ad['taches']) for ad in acts_fonct)

    suivi_map = _load_suivis_global(annee.id)

    # ── Filtre par nature d'activité ─────────────────────────────────────────
    def _filter_nature(src, invest_only):
        """Conserve uniquement les activités d'une nature donnée."""
        result = []
        for pd in src:
            pjd_list = []
            for pjd in pd['projets']:
                if invest_only:
                    acts = [ad for ad in pjd['activites']
                            if (ad['activite'].type_activite or '') == INVEST]
                else:
                    acts = [ad for ad in pjd['activites']
                            if (ad['activite'].type_activite or '') != INVEST]
                if acts:
                    pjd_copy = dict(pjd)
                    pjd_copy['activites'] = acts
                    pjd_list.append(pjd_copy)
            if pjd_list:
                pd_copy = dict(pd)
                pd_copy['projets'] = pjd_list
                result.append(pd_copy)
        return result

    # ── Calcul taux + comptages par trimestre sur un sous-ensemble ───────────
    def _stats_par_trim(data_brut):
        res = {}
        for trim in (1, 2, 3, 4, 0):
            data_f = _filter_and_renorm(data_brut, trim)
            if not data_f:
                res[trim] = {
                    'taux': 0.0,
                    'taches':    {'execute': 0, 'en_cours': 0, 'non_execute': 0, 'total': 0},
                    'activites': {'execute': 0, 'en_cours': 0, 'non_execute': 0, 'total': 0},
                }
                continue
            taux_gl, _ = _enrich(data_f, suivi_map, None)
            all_t = [td for pd in data_f for pjd in pd['projets']
                     for ad in pjd['activites'] for td in ad['taches']]
            all_a = [ad for pd in data_f for pjd in pd['projets']
                     for ad in pjd['activites']]
            res[trim] = {
                'taux': taux_gl,
                'taches': {
                    'execute':     sum(1 for td in all_t if td.get('statut') == 'execute'),
                    'en_cours':    sum(1 for td in all_t if td.get('statut') == 'en_cours'),
                    'non_execute': sum(1 for td in all_t if td.get('statut') == 'non_execute'),
                    'total': len(all_t),
                },
                'activites': {
                    'execute':     sum(1 for ad in all_a if ad.get('statut') == 'execute'),
                    'en_cours':    sum(1 for ad in all_a if ad.get('statut') == 'en_cours'),
                    'non_execute': sum(1 for ad in all_a if ad.get('statut') == 'non_execute'),
                    'total': len(all_a),
                },
            }
        return res

    suivi = _stats_par_trim(data)

    # Ventilation par nature : calculée seulement pour les admins
    suivi_invest = _stats_par_trim(_filter_nature(data, True))  if with_nature else None
    suivi_fonct  = _stats_par_trim(_filter_nature(data, False)) if with_nature else None

    # Cible Global (fin d'année) = 100 %
    cibles[0] = 100.0

    return {
        'nb_programmes':    nb_programmes,
        'nb_projets':       nb_projets,
        'nb_activites':     nb_activites,
        'nb_taches':        nb_taches,
        'nb_invest':        nb_invest,
        'nb_fonct':         nb_fonct,
        'nb_taches_invest': nb_taches_invest,
        'nb_taches_fonct':  nb_taches_fonct,
        'cibles':        cibles,
        'suivi':         suivi,
        'suivi_invest':  suivi_invest,
        'suivi_fonct':   suivi_fonct,
    }


# ─────────────────────────────────────────────────────────────────────────────
#  Route principale
# ─────────────────────────────────────────────────────────────────────────────

@dashboard_bp.route('/')
@login_required
def index():
    annee = _get_annee()
    if not annee:
        flash("Aucune année active.", 'danger')
        return redirect(url_for('pta.global_pta'))

    role         = current_user.role
    services_dir = []
    all_directions = []
    all_services   = []
    sel_svc_id   = None
    sel_dir_id   = None
    titre        = ''
    entite       = None
    data         = []
    cibles       = {1: 0.0, 2: 0.0, 3: 0.0, 4: 0.0}

    is_admin = role in ('admin_editeur', 'admin_lecteur')

    # ── Cas : service ────────────────────────────────────────────────────────
    if role == 'service':
        service = current_user.service
        if not service:
            flash("Aucun service lié à votre compte.", 'danger')
            return redirect(url_for('suivi.index'))
        from svcpta.routes import _compute_pta_service
        data   = _compute_pta_service(annee, service)
        titre  = f"{service.code} — {service.nom}"
        entite = service
        cibles = _cibles_service(annee, service)

    # ── Cas : direction ──────────────────────────────────────────────────────
    elif role == 'direction':
        direction = current_user.direction
        if not direction:
            flash("Aucune direction liée à votre compte.", 'danger')
            return redirect(url_for('suivi.index'))
        services_dir = Service.query.filter_by(direction_id=direction.id)\
                                    .order_by(Service.nom).all()
        sel_svc_id = request.args.get('service_id', type=int)
        if sel_svc_id and not any(s.id == sel_svc_id for s in services_dir):
            sel_svc_id = None

        if sel_svc_id:
            svc = Service.query.get(sel_svc_id)
            from svcpta.routes import _compute_pta_service
            data   = _compute_pta_service(annee, svc)
            titre  = f"{svc.code} — {svc.nom}"
            entite = svc
            cibles = _cibles_service(annee, svc)
        else:
            from dirpta.routes import _compute_pta_direction
            data   = _compute_pta_direction(annee, direction)
            titre  = f"{direction.code} — {direction.nom}"
            entite = direction
            cibles = _cibles_direction(annee, direction)

    # ── Cas : admin (éditeur ou lecteur) ────────────────────────────────────
    elif is_admin:
        all_directions = Direction.query.order_by(Direction.nom).all()
        all_services   = Service.query.order_by(Service.nom).all()

        sel_dir_id = request.args.get('direction_id', type=int)
        sel_svc_id = request.args.get('service_id', type=int)

        # Priorité : filtre service > filtre direction > global
        if sel_svc_id:
            svc = Service.query.get(sel_svc_id)
            if not svc:
                sel_svc_id = None
            else:
                sel_dir_id = None
                from svcpta.routes import _compute_pta_service
                data   = _compute_pta_service(annee, svc)
                titre  = f"Service — {svc.code} · {svc.nom}"
                entite = svc
                cibles = _cibles_service(annee, svc)

        if not sel_svc_id and sel_dir_id:
            direction = Direction.query.get(sel_dir_id)
            if not direction:
                sel_dir_id = None
            else:
                from dirpta.routes import _compute_pta_direction
                data   = _compute_pta_direction(annee, direction)
                titre  = f"Direction — {direction.code} · {direction.nom}"
                entite = direction
                cibles = _cibles_direction(annee, direction)

        if not sel_svc_id and not sel_dir_id:
            from suivi.routes import _compute_pta_global
            data   = _compute_pta_global(annee)
            titre  = "PTA Global — Mairie d'Adja-Ouèrè"
            entite = None
            cibles = _cibles_global(annee)

    stats    = _compute_dashboard_stats(annee, data, cibles,
                                        with_nature=is_admin) if data else None
    synthese = _compute_synthese_admin(annee) if is_admin else None

    return render_template('dashboard/index.html',
        annee=annee,
        titre=titre,
        role=role,
        entite=entite,
        is_admin=is_admin,
        services_dir=services_dir,
        sel_svc_id=sel_svc_id,
        all_directions=all_directions,
        all_services=all_services,
        sel_dir_id=sel_dir_id,
        stats=stats,
        synthese=synthese,
    )


# ─────────────────────────────────────────────────────────────────────────────
#  Construction du document Word
# ─────────────────────────────────────────────────────────────────────────────

def _build_dashboard_word(annee, titre, stats, synthese, static_img_path):
    """Génère un document Word commenté du tableau de bord admin."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── Helpers XML ──────────────────────────────────────────────────────────

    def shade(cell, hex_col):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_col)
        tcPr.append(shd)

    def borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBrd = OxmlElement('w:tblBorders')
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'AAAAAA')
            tblBrd.append(b)
        tblPr.append(tblBrd)

    def repeat_hdr(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        hdr = OxmlElement('w:tblHeader')
        hdr.set(qn('w:val'), 'true')
        trPr.append(hdr)

    # ── Helpers texte ────────────────────────────────────────────────────────

    def h1(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    def h2(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

    def interp(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run('ℹ  ' + text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    def hdr_cell(cell, text, bg='1F4E79'):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, bg)

    def data_cell(cell, text, align='left', bold=False, size=9, fg=None, bg=None):
        p = cell.paragraphs[0]
        p.alignment = {'left':   WD_ALIGN_PARAGRAPH.LEFT,
                       'center': WD_ALIGN_PARAGRAPH.CENTER,
                       'right':  WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(str(text) if text is not None else '')
        r.bold = bold
        r.font.size = Pt(size)
        if fg:
            r.font.color.rgb = RGBColor(*fg)
        if bg:
            shade(cell, bg)

    def taux_color(taux):
        """Couleur RGB selon le niveau de taux."""
        if taux >= 75:
            return (0x15, 0x57, 0x24)   # vert foncé
        elif taux >= 40:
            return (0x85, 0x64, 0x04)   # orange foncé
        else:
            return (0xDC, 0x35, 0x45)   # rouge

    def taux_bg(taux):
        """Fond cellule selon le niveau de taux."""
        if taux >= 75:
            return 'D5F5E3'
        elif taux >= 40:
            return 'FFF3CD'
        else:
            return 'FADBD8'

    # ── Tableau statuts (tâches ou activités) ────────────────────────────────

    def table_statuts(doc, suivi_data, section):
        """Tableau T1/T2/T3/T4/Global des statuts (tâches ou activités)."""
        TRIMS = (1, 2, 3, 4, 0)
        tbl = doc.add_table(rows=5, cols=6)
        borders(tbl)
        hdr_cell(tbl.cell(0, 0), 'Statut', bg='1F4E79')
        for ci, lbl in enumerate(['T1', 'T2', 'T3', 'T4', 'Global'], 1):
            bg = '155724' if ci == 5 else '1F4E79'
            hdr_cell(tbl.cell(0, ci), lbl, bg=bg)
        repeat_hdr(tbl.rows[0])

        STATUTS = [
            ('execute',     'Exécutées',      '198754'),
            ('en_cours',    'En cours',       'C77700'),
            ('non_execute', 'Non exécutées',  '6C757D'),
        ]
        for ri, (skey, slbl, col) in enumerate(STATUTS, 1):
            r = tbl.rows[ri]
            data_cell(r.cells[0], slbl, bold=True, size=9)
            shade(r.cells[0], 'F8F9FA')
            for ci, t in enumerate(TRIMS, 1):
                grp = suivi_data[t]['taches'] if section == 'taches' else suivi_data[t]['activites']
                n   = grp[skey]
                tot = grp['total']
                txt = f"{n}  ({n*100//tot} %)" if tot > 0 else "—"
                bg_c = 'D1F0DA' if ci == 5 else 'FFFFFF'
                data_cell(r.cells[ci], txt, align='center', size=9, bg=bg_c)

        # Ligne total
        r = tbl.rows[4]
        data_cell(r.cells[0], 'Total', bold=True, size=9, bg='E8EAF6')
        for ci, t in enumerate(TRIMS, 1):
            grp = suivi_data[t]['taches'] if section == 'taches' else suivi_data[t]['activites']
            bg_c = 'D1F0DA' if ci == 5 else 'E8EAF6'
            data_cell(r.cells[ci], str(grp['total']), align='center', bold=True, size=9, bg=bg_c)
        return tbl

    # ── Commentaire automatique statuts ──────────────────────────────────────

    def commentaire_statuts(suivi_data, section, label):
        grp = suivi_data[0]['taches'] if section == 'taches' else suivi_data[0]['activites']
        tot = grp['total']
        if tot == 0:
            return f"Aucune {label} enregistrée."
        ex  = grp['execute']
        ec  = grp['en_cours']
        ne  = grp['non_execute']
        p_ex = round(ex * 100 / tot, 1)
        p_ec = round(ec * 100 / tot, 1)
        p_ne = round(ne * 100 / tot, 1)
        txt = (f"Sur {tot} {label}, {ex} sont exécutées ({p_ex} %), "
               f"{ec} sont en cours d'exécution ({p_ec} %) "
               f"et {ne} n'ont pas encore démarré ({p_ne} %). ")
        if p_ex >= 75:
            txt += "Le niveau d'exécution est très satisfaisant."
        elif p_ex >= 50:
            txt += "Plus de la moitié des tâches sont achevées, les efforts doivent se poursuivre."
        elif p_ne >= 50:
            txt += "Plus de la moitié des tâches restent à exécuter ; une accélération s'impose."
        else:
            txt += "L'exécution est en progression mais des efforts supplémentaires sont nécessaires."
        return txt

    # ── Construction du document ─────────────────────────────────────────────

    doc = Document()
    sec0 = doc.sections[0]
    sec0.orientation   = WD_ORIENT.PORTRAIT
    sec0.page_width    = Cm(21)
    sec0.page_height   = Cm(29.7)
    sec0.left_margin   = sec0.right_margin  = Cm(2.5)
    sec0.top_margin    = Cm(2.5)
    sec0.bottom_margin = Cm(2)

    # ── Page de titre ────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('TABLEAU DE BORD')
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'PLAN DE TRAVAIL ANNUEL — Exercice {annee.annee}')
    r2.bold = True; r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(titre)
    r3.bold = True; r3.font.size = Pt(12)

    logo_path = os.path.join(static_img_path, 'logo_commune.png')
    if os.path.exists(logo_path):
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.add_run().add_picture(logo_path, height=Cm(2.4))

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')
    r4 = p4.add_run(f'Généré le {date_str}')
    r4.italic = True; r4.font.size = Pt(9)
    r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.add_run('─' * 55).font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    if not stats:
        p = doc.add_paragraph()
        p.add_run("Aucune donnée disponible pour cette entité.").italic = True
        return doc

    # ════════════════════════════════════════════════════════════════════════
    # I. STRUCTURE DU PTA
    # ════════════════════════════════════════════════════════════════════════
    h1(doc, 'I. STRUCTURE DU PTA')

    # Tableau KPI (4 colonnes : Programmes / Projets / Activités / Tâches)
    tbl_kpi = doc.add_table(rows=2, cols=4)
    borders(tbl_kpi)
    for ci, (lbl, val) in enumerate([
        ('Programmes', stats['nb_programmes']),
        ('Projets',    stats['nb_projets']),
        ('Activités',  stats['nb_activites']),
        ('Tâches',     stats['nb_taches']),
    ]):
        hdr_cell(tbl_kpi.cell(0, ci), lbl)
        c = tbl_kpi.cell(1, ci)
        data_cell(c, str(val), align='center', bold=True, size=12, bg='D6EAF8')
    repeat_hdr(tbl_kpi.rows[0])
    doc.add_paragraph()

    # Tableau répartition par nature
    h2(doc, "Répartition par nature d'activité")
    tbl_nat = doc.add_table(rows=3, cols=3)
    borders(tbl_nat)
    for ci, lbl in enumerate(["Nature d'activité", 'Activités', 'Tâches']):
        hdr_cell(tbl_nat.cell(0, ci), lbl, bg='2C3E50')
    repeat_hdr(tbl_nat.rows[0])
    for ri, (lbl, nb_a, nb_t, bg_c) in enumerate([
        ("Investissement",  stats['nb_invest'], stats['nb_taches_invest'], 'D6EAF8'),
        ("Fonctionnement",  stats['nb_fonct'],  stats['nb_taches_fonct'],  'FEF9E7'),
    ], 1):
        r = tbl_nat.rows[ri]
        data_cell(r.cells[0], lbl, bold=True, size=9, bg=bg_c)
        data_cell(r.cells[1], str(nb_a), align='center', size=9, bg=bg_c)
        data_cell(r.cells[2], str(nb_t), align='center', size=9, bg=bg_c)

    total_a = stats['nb_activites']
    pct_inv = round(stats['nb_invest'] * 100 / total_a, 1) if total_a else 0
    interp(doc, (
        f"Le PTA comprend {stats['nb_activites']} activités pour {stats['nb_taches']} tâches, "
        f"dont {stats['nb_invest']} activités d'investissement ({pct_inv} %) "
        f"et {stats['nb_fonct']} de fonctionnement."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # II. TAUX D'EXÉCUTION
    # ════════════════════════════════════════════════════════════════════════
    h1(doc, "II. TAUX D'EXÉCUTION")

    tbl_tx = doc.add_table(rows=3, cols=6)
    borders(tbl_tx)
    hdr_cell(tbl_tx.cell(0, 0), 'Indicateur', bg='1F4E79')
    for ci, lbl in enumerate(['T1', 'T2', 'T3', 'T4'], 1):
        hdr_cell(tbl_tx.cell(0, ci), lbl, bg='1F4E79')
    hdr_cell(tbl_tx.cell(0, 5), 'Global', bg='155724')
    repeat_hdr(tbl_tx.rows[0])

    # Ligne cibles
    r_cib = tbl_tx.rows[1]
    data_cell(r_cib.cells[0], 'Cible théorique', bold=True, size=9, bg='F8F9FA')
    for ci, t in enumerate((1, 2, 3, 4), 1):
        v = stats['cibles'].get(t, 0)
        bg_c = 'D4E6F1' if v >= 50 else 'FDEBD0'
        data_cell(r_cib.cells[ci], f"{v:.1f} %", align='center', bold=True, size=9,
                  fg=(0x1F, 0x4E, 0x79), bg=bg_c)
    data_cell(r_cib.cells[5], f"{stats['cibles'][0]:.1f} %",
              align='center', bold=True, size=9, fg=(0x1F, 0x4E, 0x79), bg='D1F0DA')

    # Ligne taux réel (uniquement Global)
    r_tx = tbl_tx.rows[2]
    data_cell(r_tx.cells[0], 'Taux réel atteint', bold=True, size=9, bg='F8F9FA')
    for ci in range(1, 5):
        data_cell(r_tx.cells[ci], '—', align='center', size=9, bg='F8F9FA')
    taux_gl = stats['suivi'][0]['taux']
    data_cell(r_tx.cells[5], f"{taux_gl:.1f} %",
              align='center', bold=True, size=9,
              fg=taux_color(taux_gl), bg='D1F0DA')

    ecart = taux_gl - stats['cibles'][0]
    signe = '+' if ecart >= 0 else ''
    if ecart >= 0:
        cmt_tx = (f"Le taux de réalisation global s'établit à {taux_gl:.1f} %, "
                  f"soit {signe}{ecart:.1f} point(s) par rapport à la cible de 100 %. "
                  f"L'entité dépasse ses objectifs.")
    elif ecart >= -20:
        cmt_tx = (f"Le taux de réalisation global atteint {taux_gl:.1f} %, "
                  f"accusant un retard de {abs(ecart):.1f} point(s) par rapport à la cible de 100 %. "
                  f"Des efforts soutenus sont nécessaires pour combler l'écart.")
    else:
        cmt_tx = (f"Le taux de réalisation global est de {taux_gl:.1f} %, "
                  f"soit un retard significatif de {abs(ecart):.1f} point(s) par rapport à la cible de 100 %. "
                  f"Une revue des facteurs bloquants et une accélération de l'exécution s'imposent.")
    doc.add_paragraph()
    interp(doc, cmt_tx)

    # ════════════════════════════════════════════════════════════════════════
    # III. TÂCHES PAR STATUT
    # ════════════════════════════════════════════════════════════════════════
    h1(doc, 'III. TÂCHES PAR STATUT')

    h2(doc, 'Toutes natures confondues')
    table_statuts(doc, stats['suivi'], 'taches')
    doc.add_paragraph()
    interp(doc, commentaire_statuts(stats['suivi'], 'taches', 'tâches'))

    if stats.get('suivi_invest'):
        h2(doc, "Activités d'investissement")
        table_statuts(doc, stats['suivi_invest'], 'taches')
        doc.add_paragraph()
        interp(doc, commentaire_statuts(stats['suivi_invest'], 'taches',
                                        "tâches d'investissement"))

    if stats.get('suivi_fonct'):
        h2(doc, 'Activités de fonctionnement')
        table_statuts(doc, stats['suivi_fonct'], 'taches')
        doc.add_paragraph()
        interp(doc, commentaire_statuts(stats['suivi_fonct'], 'taches',
                                        'tâches de fonctionnement'))

    # ════════════════════════════════════════════════════════════════════════
    # IV. ACTIVITÉS PAR STATUT
    # ════════════════════════════════════════════════════════════════════════
    h1(doc, 'IV. ACTIVITÉS PAR STATUT')

    h2(doc, 'Toutes natures confondues')
    table_statuts(doc, stats['suivi'], 'activites')
    doc.add_paragraph()
    interp(doc, commentaire_statuts(stats['suivi'], 'activites', 'activités'))

    if stats.get('suivi_invest'):
        h2(doc, "Investissement")
        table_statuts(doc, stats['suivi_invest'], 'activites')
        doc.add_paragraph()
        interp(doc, commentaire_statuts(stats['suivi_invest'], 'activites',
                                        "activités d'investissement"))

    if stats.get('suivi_fonct'):
        h2(doc, 'Fonctionnement')
        table_statuts(doc, stats['suivi_fonct'], 'activites')
        doc.add_paragraph()
        interp(doc, commentaire_statuts(stats['suivi_fonct'], 'activites',
                                        'activités de fonctionnement'))

    # ════════════════════════════════════════════════════════════════════════
    # V. VUE D'ENSEMBLE — DIRECTIONS & SERVICES
    # ════════════════════════════════════════════════════════════════════════
    if synthese:
        h1(doc, "V. VUE D'ENSEMBLE — DIRECTIONS & SERVICES")

        tbl_syn = doc.add_table(rows=1, cols=2)
        borders(tbl_syn)
        hdr_cell(tbl_syn.cell(0, 0), 'Entité', bg='1F4E79')
        hdr_cell(tbl_syn.cell(0, 1), 'Taux global (%)', bg='1F4E79')
        repeat_hdr(tbl_syn.rows[0])

        taux_items = []   # pour commentaire final

        for item in synthese:
            td = item['taux']
            taux_items.append((item['direction'].nom, td))
            r = tbl_syn.add_row()
            data_cell(r.cells[0],
                      f"{item['direction'].code} — {item['direction'].nom}",
                      bold=True, size=9, bg='E8EAF6')
            data_cell(r.cells[1], f"{td:.1f} %",
                      align='center', bold=True, size=9,
                      fg=taux_color(td), bg=taux_bg(td))

            for sv in item['services']:
                ts = sv['taux']
                taux_items.append((f"  └ {sv['service'].nom}", ts))
                rs = tbl_syn.add_row()
                data_cell(rs.cells[0],
                          f"   {sv['service'].code} — {sv['service'].nom}",
                          size=9, bg='FDFEFE')
                data_cell(rs.cells[1], f"{ts:.1f} %",
                          align='center', size=9,
                          fg=taux_color(ts), bg=taux_bg(ts))

        doc.add_paragraph()

        # Commentaire : meilleure et moins bonne direction
        dirs_only = [(item['direction'].nom, item['taux']) for item in synthese]
        if dirs_only:
            best = max(dirs_only, key=lambda x: x[1])
            worst = min(dirs_only, key=lambda x: x[1])
            if best[0] == worst[0]:
                interp(doc, (f"La direction {best[0]} affiche un taux de {best[1]:.1f} %."))
            else:
                interp(doc, (
                    f"La direction la plus avancée est {best[0]} avec {best[1]:.1f} %. "
                    f"La direction {worst[0]} affiche le taux le plus bas ({worst[1]:.1f} %) "
                    f"et mérite une attention particulière."
                ))

    return doc


# ─────────────────────────────────────────────────────────────────────────────
#  Route export Word
# ─────────────────────────────────────────────────────────────────────────────

@dashboard_bp.route('/rapport/word')
@login_required
def rapport_word():
    """Export Word du tableau de bord — réservé aux admins."""
    if current_user.role not in ('admin_editeur', 'admin_lecteur'):
        flash("Accès réservé aux administrateurs.", 'warning')
        return redirect(url_for('dashboard.index'))

    annee = _get_annee()
    if not annee:
        flash("Aucune année active.", 'danger')
        return redirect(url_for('dashboard.index'))

    try:
        sel_dir_id = request.args.get('direction_id', type=int)
        sel_svc_id = request.args.get('service_id', type=int)
        cibles = {1: 0.0, 2: 0.0, 3: 0.0, 4: 0.0}
        titre  = ''
        data   = []

        if sel_svc_id:
            svc = Service.query.get(sel_svc_id)
            if svc:
                sel_dir_id = None
                from svcpta.routes import _compute_pta_service
                data   = _compute_pta_service(annee, svc)
                titre  = f"Service — {svc.code} · {svc.nom}"
                cibles = _cibles_service(annee, svc)

        if not sel_svc_id and sel_dir_id:
            direction = Direction.query.get(sel_dir_id)
            if direction:
                from dirpta.routes import _compute_pta_direction
                data   = _compute_pta_direction(annee, direction)
                titre  = f"Direction — {direction.code} · {direction.nom}"
                cibles = _cibles_direction(annee, direction)

        if not sel_svc_id and not sel_dir_id:
            from suivi.routes import _compute_pta_global
            data   = _compute_pta_global(annee)
            titre  = "PTA Global — Mairie d'Adja-Ouèrè"
            cibles = _cibles_global(annee)

        stats    = _compute_dashboard_stats(annee, data, cibles, with_nature=True) if data else None
        synthese = _compute_synthese_admin(annee)

        static_img = os.path.join(current_app.root_path, 'static', 'img')
        doc = _build_dashboard_word(annee, titre, stats, synthese, static_img)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        # Nom de fichier propre
        safe = titre.replace(' ', '_').replace('—', '-').replace('·', '').replace("'", '')
        safe = ''.join(c for c in safe if c.isalnum() or c in ('_', '-'))
        fname = f"Tableau_Bord_{safe}_{annee.annee}.docx"

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=fname,
        )

    except Exception as e:
        flash(f'Erreur lors de la génération du rapport : {e}', 'danger')
        return redirect(url_for('dashboard.index'))
