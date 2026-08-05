import io
import os
import datetime
from flask import render_template, session, redirect, url_for, flash, send_file, current_app
from flask_login import login_required
from sqlalchemy import func
from models import db, Programme, Projet, Activite, Tache, Direction, Service, Annee
from stats import stats_bp

INVEST = "Activité d'investissement"
FONCT  = "Activité de fonctionnement"

MOIS_NUM = {
    'Janvier':1,'Février':2,'Mars':3,'Avril':4,'Mai':5,'Juin':6,
    'Juillet':7,'Août':8,'Septembre':9,'Octobre':10,'Novembre':11,'Décembre':12,
}
END_TRIM = {1: 3, 2: 6, 3: 9, 4: 12}


# ── Helpers ───────────────────────────────────────────────────────────────────

def get_annee():
    annee_id = session.get('annee_id')
    if annee_id:
        return Annee.query.get(annee_id)
    return Annee.query.filter_by(actif=True).first()


def _fmt(v):
    try:
        return '{:,.0f}'.format(float(v or 0)).replace(',', ' ')
    except (TypeError, ValueError):
        return '0'


def _compute_stats(annee):
    programmes = Programme.query.filter_by(annee_id=annee.id).order_by(Programme.numero).all()
    directions = Direction.query.order_by(Direction.nom).all()

    all_projets   = [p for prog in programmes for p in prog.projets]
    all_activites = [a for p in all_projets   for a in p.activites]
    all_taches    = [t for a in all_activites  for t in a.taches]

    all_activites_invest = [a for a in all_activites if a.type_activite == INVEST]
    all_activites_fonct  = [a for a in all_activites if a.type_activite != INVEST]
    all_taches_invest    = [t for a in all_activites_invest for t in a.taches]
    all_taches_fonct     = [t for a in all_activites_fonct  for t in a.taches]

    def _src(acts):
        rp = sum(a.src_rp for a in acts)
        fa = sum(a.src_fa for a in acts)
        fn = sum(a.src_fn for a in acts)
        ap = sum(a.src_ap for a in acts)
        af = sum(a.src_af for a in acts)
        return rp, fa, fn, ap, af, rp + fa + fn + ap + af

    total_rp, total_fa, total_fn, total_ap, total_af, total_budget = _src(all_activites)
    inv_rp,   inv_fa,   inv_fn,   inv_ap,   inv_af,   inv_budget   = _src(all_activites_invest)
    fct_rp,   fct_fa,   fct_fn,   fct_ap,   fct_af,   fct_budget   = _src(all_activites_fonct)

    def pct(val):        return round(val / total_budget * 100, 1) if total_budget else 0
    def pct_invest(val): return round(val / inv_budget   * 100, 1) if inv_budget   else 0
    def pct_fonct(val):  return round(val / fct_budget   * 100, 1) if fct_budget   else 0

    def _proj_entry(proj, type_filter=None):
        acts = [a for a in proj.activites
                if type_filter is None or a.type_activite == type_filter]
        taches = [t for a in acts for t in a.taches]
        rp, fa, fn, ap, af, budget = _src(acts)
        return {
            'code': proj.code, 'nom': proj.nom, 'poids': proj.poids or 0,
            'nb_activites': len(acts), 'nb_taches': len(taches),
            'nb_activites_invest': len([a for a in proj.activites if a.type_activite == INVEST]),
            'nb_activites_fonct':  len([a for a in proj.activites if a.type_activite != INVEST]),
            'budget': budget, 'rp': rp, 'fa': fa, 'fn': fn, 'ap': ap, 'af': af,
        }

    def _prog_entry(prog, type_filter=None):
        projets_data = []
        for proj in prog.projets:
            e = _proj_entry(proj, type_filter)
            if type_filter is None or e['nb_activites'] > 0:
                projets_data.append(e)
        acts = [a for p in prog.projets for a in p.activites
                if type_filter is None or a.type_activite == type_filter]
        taches = [t for a in acts for t in a.taches]
        rp, fa, fn, ap, af, budget = _src(acts)
        return {
            'code': prog.code, 'nom': prog.nom, 'poids': prog.poids or 0,
            'nb_projets': len([p for p in prog.projets
                               if any(type_filter is None or a.type_activite == type_filter
                                      for a in p.activites)]),
            'nb_activites': len(acts), 'nb_taches': len(taches),
            'nb_activites_invest': len([a for p in prog.projets for a in p.activites if a.type_activite == INVEST]),
            'nb_activites_fonct':  len([a for p in prog.projets for a in p.activites if a.type_activite != INVEST]),
            'budget': budget, 'rp': rp, 'fa': fa, 'fn': fn, 'ap': ap, 'af': af,
            'projets': projets_data,
        }

    stats_par_programme        = [_prog_entry(p)       for p in programmes]
    stats_par_programme_invest = [e for e in [_prog_entry(p, INVEST) for p in programmes] if e['nb_activites'] > 0]
    stats_par_programme_fonct  = [e for e in [_prog_entry(p, FONCT)  for p in programmes] if e['nb_activites'] > 0]

    def _dir_stats(type_filter=None):
        q_act = (db.session.query(Activite.direction_responsable_id, func.count(Activite.id))
                 .join(Projet,    Projet.id    == Activite.projet_id)
                 .join(Programme, Programme.id == Projet.programme_id)
                 .filter(Programme.annee_id == annee.id))
        q_tch = (db.session.query(Tache.direction_responsable_id, func.count(Tache.id))
                 .join(Activite,  Activite.id  == Tache.activite_id)
                 .join(Projet,    Projet.id    == Activite.projet_id)
                 .join(Programme, Programme.id == Projet.programme_id)
                 .filter(Programme.annee_id == annee.id))
        if type_filter:
            q_act = q_act.filter(Activite.type_activite == type_filter)
            q_tch = q_tch.filter(Activite.type_activite == type_filter)
        act_map = {r[0]: r[1] for r in q_act.group_by(Activite.direction_responsable_id).all()}
        tch_map = {r[0]: r[1] for r in q_tch.group_by(Tache.direction_responsable_id).all()}
        return [{'code': d.code, 'nom': d.nom,
                 'nb_activites': act_map.get(d.id, 0),
                 'nb_taches':    tch_map.get(d.id, 0)} for d in directions]

    stats_directions        = _dir_stats()
    stats_directions_invest = _dir_stats(INVEST)
    stats_directions_fonct  = _dir_stats(FONCT)

    # Liste détaillée des activités d'investissement
    activites_investissement = []
    for a in sorted(all_activites_invest,
                    key=lambda x: (x.projet.programme.numero, x.projet.numero, x.numero)):
        budget_a = a.src_rp + a.src_fa + a.src_fn + a.src_ap + a.src_af
        activites_investissement.append({
            'code': a.code, 'nom': a.nom,
            'direction': a.direction_responsable.code if a.direction_responsable else '—',
            'direction_nom': a.direction_responsable.nom if a.direction_responsable else '—',
            'periode': (f"{a.periode_debut}–{a.periode_fin}" if a.periode_debut and a.periode_fin
                        else a.periode_debut or a.periode_fin or '—'),
            'poids': a.poids or 0,
            'rp': a.src_rp, 'fa': a.src_fa, 'fn': a.src_fn, 'ap': a.src_ap, 'af': a.src_af,
            'budget': budget_a,
            'nb_taches': len(a.taches),
        })

    # Liste détaillée des activités de fonctionnement
    activites_fonctionnement = []
    for a in sorted(all_activites_fonct,
                    key=lambda x: (x.projet.programme.numero, x.projet.numero, x.numero)):
        budget_a = a.src_rp + a.src_fa + a.src_fn + a.src_ap + a.src_af
        activites_fonctionnement.append({
            'code': a.code, 'nom': a.nom,
            'direction': a.direction_responsable.code if a.direction_responsable else '—',
            'direction_nom': a.direction_responsable.nom if a.direction_responsable else '—',
            'periode': (f"{a.periode_debut}–{a.periode_fin}" if a.periode_debut and a.periode_fin
                        else a.periode_debut or a.periode_fin or '—'),
            'poids': a.poids or 0,
            'rp': a.src_rp, 'fa': a.src_fa, 'fn': a.src_fn, 'ap': a.src_ap, 'af': a.src_af,
            'budget': budget_a,
            'nb_taches': len(a.taches),
        })

    return dict(
        programmes=programmes, directions=directions,
        total_budget=total_budget,
        total_rp=total_rp, total_fa=total_fa, total_fn=total_fn,
        total_ap=total_ap, total_af=total_af, pct=pct,
        nb_programmes=len(programmes), nb_projets=len(all_projets),
        nb_activites=len(all_activites), nb_taches=len(all_taches),
        nb_activites_invest=len(all_activites_invest),
        nb_activites_fonct=len(all_activites_fonct),
        nb_taches_invest=len(all_taches_invest),
        nb_taches_fonct=len(all_taches_fonct),
        inv_rp=inv_rp, inv_fa=inv_fa, inv_fn=inv_fn, inv_ap=inv_ap, inv_af=inv_af,
        inv_budget=inv_budget, pct_invest=pct_invest,
        fct_rp=fct_rp, fct_fa=fct_fa, fct_fn=fct_fn, fct_ap=fct_ap, fct_af=fct_af,
        fct_budget=fct_budget, pct_fonct=pct_fonct,
        stats_par_programme=stats_par_programme,
        stats_par_programme_invest=stats_par_programme_invest,
        stats_par_programme_fonct=stats_par_programme_fonct,
        stats_directions=stats_directions,
        stats_directions_invest=stats_directions_invest,
        stats_directions_fonct=stats_directions_fonct,
        activites_investissement=activites_investissement,
        activites_fonctionnement=activites_fonctionnement,
    )


# ── Cibles trimestrielles ─────────────────────────────────────────────────────

def _tache_contrib(t, trim):
    """Proportion (0.0–1.0) d'une tâche réalisée au cumul à fin du trimestre `trim`.
    Règle : progression linéaire en mois sur l'intervalle [periode_debut, periode_fin].
    Si la période n'est pas renseignée, on suppose Janvier→Décembre (25%/50%/75%/100%).
    """
    m_d = MOIS_NUM.get(t.periode_debut or '', 1)
    m_f = MOIS_NUM.get(t.periode_fin   or '', 12)
    if m_f < m_d:
        m_f = m_d
    duree     = m_f - m_d + 1
    completed = min(m_f, END_TRIM[trim]) - m_d + 1
    return max(0.0, min(completed / duree, 1.0))


def _compute_cibles_directions(annee, directions):
    """Calcule pour chaque direction la cible trimestrielle globale.

    Utilise directement _compute_pta_direction (module dirpta) qui fournit
    la hiérarchie déjà filtrée avec les poids renormalisés (new_poids).
    Applique ensuite la même remontée hiérarchique que _compute_cibles :
    tâche → activité → projet → programme.
    """
    from dirpta.routes import _compute_pta_direction
    result = []

    for direction in directions:
        data = _compute_pta_direction(annee, direction)
        if not data:
            continue

        prog_entries = []   # [(new_poids_prog, {trim: cible}), …]
        nb_taches_dir = 0

        for pd in data:                        # pd  = dict programme
            proj_entries = []

            for pjd in pd['projets']:          # pjd = dict projet
                act_entries = []

                for ad in pjd['activites']:    # ad  = dict activité
                    taches = ad['taches']      # liste de dicts tâche
                    if not taches:
                        continue
                    tot_tp = sum(td['new_poids'] for td in taches)
                    if tot_tp == 0:
                        continue
                    cibles_act = {
                        trim: sum(_tache_contrib(td['tache'], trim) * td['new_poids']
                                  for td in taches) / tot_tp * 100
                        for trim in (1, 2, 3, 4)
                    }
                    act_entries.append((ad['new_poids'], cibles_act))
                    nb_taches_dir += len(taches)

                if not act_entries:
                    continue
                tot_ap = sum(p for p, _ in act_entries)
                if tot_ap == 0:
                    continue
                cibles_proj = {
                    trim: sum(p * c[trim] for p, c in act_entries) / tot_ap
                    for trim in (1, 2, 3, 4)
                }
                proj_entries.append((pjd['new_poids'], cibles_proj))

            if not proj_entries:
                continue
            tot_pp = sum(p for p, _ in proj_entries)
            if tot_pp == 0:
                continue
            cibles_prog = {
                trim: sum(p * c[trim] for p, c in proj_entries) / tot_pp
                for trim in (1, 2, 3, 4)
            }
            prog_entries.append((pd['new_poids'], cibles_prog))

        if not prog_entries:
            continue
        tot_pgp = sum(p for p, _ in prog_entries)
        if tot_pgp == 0:
            continue
        cibles_dir = {
            trim: round(sum(p * c[trim] for p, c in prog_entries) / tot_pgp, 1)
            for trim in (1, 2, 3, 4)
        }
        result.append({
            'code':      direction.code,
            'nom':       direction.nom,
            'nb_taches': nb_taches_dir,
            'cibles':    cibles_dir,
        })

    return result


def _compute_cibles_services(annee, services):
    """Calcule pour chaque service la cible trimestrielle globale.

    Même logique que _compute_cibles_directions mais en utilisant
    _compute_pta_service (module svcpta).
    """
    from svcpta.routes import _compute_pta_service
    result = []

    for service in services:
        data = _compute_pta_service(annee, service)
        if not data:
            continue

        prog_entries = []   # [(new_poids_prog, {trim: cible}), …]
        nb_taches_svc = 0

        for pd in data:                        # pd  = dict programme
            proj_entries = []

            for pjd in pd['projets']:          # pjd = dict projet
                act_entries = []

                for ad in pjd['activites']:    # ad  = dict activité
                    taches = ad['taches']      # liste de dicts tâche
                    if not taches:
                        continue
                    tot_tp = sum(td['new_poids'] for td in taches)
                    if tot_tp == 0:
                        continue
                    cibles_act = {
                        trim: sum(_tache_contrib(td['tache'], trim) * td['new_poids']
                                  for td in taches) / tot_tp * 100
                        for trim in (1, 2, 3, 4)
                    }
                    act_entries.append((ad['new_poids'], cibles_act))
                    nb_taches_svc += len(taches)

                if not act_entries:
                    continue
                tot_ap = sum(p for p, _ in act_entries)
                if tot_ap == 0:
                    continue
                cibles_proj = {
                    trim: sum(p * c[trim] for p, c in act_entries) / tot_ap
                    for trim in (1, 2, 3, 4)
                }
                proj_entries.append((pjd['new_poids'], cibles_proj))

            if not proj_entries:
                continue
            tot_pp = sum(p for p, _ in proj_entries)
            if tot_pp == 0:
                continue
            cibles_prog = {
                trim: sum(p * c[trim] for p, c in proj_entries) / tot_pp
                for trim in (1, 2, 3, 4)
            }
            prog_entries.append((pd['new_poids'], cibles_prog))

        if not prog_entries:
            continue
        tot_pgp = sum(p for p, _ in prog_entries)
        if tot_pgp == 0:
            continue
        cibles_svc = {
            trim: round(sum(p * c[trim] for p, c in prog_entries) / tot_pgp, 1)
            for trim in (1, 2, 3, 4)
        }
        result.append({
            'code':           service.code,
            'nom':            service.nom,
            'direction_code': service.direction.code if service.direction else '—',
            'nb_taches':      nb_taches_svc,
            'cibles':         cibles_svc,
        })

    return result


def _compute_cibles(programmes):
    """Calcule pour chaque trimestre (1-4) le taux d'exécution cible par programme et projet.
    Retourne :
      {
        'global': {1: xx.x, …, 4: xx.x},
        'programmes': [{
            'code':…, 'nom':…, 'poids':…,
            'cibles': {1:…, 4:…},
            'projets': [{'code':…, 'nom':…, 'poids':…, 'cibles':{1:…,4:…}}, …]
        }, …]
      }
    """
    all_prog = []

    for prog in programmes:
        proj_entries = []

        for proj in prog.projets:
            # Calcul des cibles de l'activité pour les 4 trimestres en un seul passage
            act_data = []  # (poids_act, {trim: cible_%})
            for act in proj.activites:
                taches = act.taches
                if not taches:
                    continue
                tot_pt = sum(t.poids or 0 for t in taches)
                if tot_pt == 0:
                    continue
                cibles_act = {
                    trim: sum(_tache_contrib(t, trim) * (t.poids or 0)
                              for t in taches) / tot_pt * 100
                    for trim in (1, 2, 3, 4)
                }
                act_data.append((act.poids or 0, cibles_act))

            if not act_data:
                continue
            tot_a = sum(p for p, _ in act_data)
            if tot_a == 0:
                continue

            proj_cibles = {
                trim: round(sum(p * c[trim] for p, c in act_data) / tot_a, 1)
                for trim in (1, 2, 3, 4)
            }
            proj_entries.append({
                'code':   proj.code,
                'nom':    proj.nom,
                'poids':  proj.poids or 0,
                'cibles': proj_cibles,
            })

        if not proj_entries:
            continue

        tot_pr = sum(p['poids'] for p in proj_entries)
        prog_cibles = {
            trim: round(
                sum(p['cibles'][trim] * p['poids'] for p in proj_entries) / tot_pr if tot_pr else 0,
                1
            )
            for trim in (1, 2, 3, 4)
        }
        all_prog.append({
            'code':    prog.code,
            'nom':     prog.nom,
            'poids':   prog.poids or 0,
            'cibles':  prog_cibles,
            'projets': proj_entries,
        })

    tot_pg = sum(p['poids'] for p in all_prog)
    global_cibles = {
        trim: round(
            sum(p['cibles'][trim] * p['poids'] for p in all_prog) / tot_pg if tot_pg else 0,
            1
        )
        for trim in (1, 2, 3, 4)
    }
    return {'global': global_cibles, 'programmes': all_prog}


# ── Génération des graphiques ─────────────────────────────────────────────────

def _chart_sources(rp, fa, fn, ap, af, title='Répartition par source de financement'):
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        labels = ['Fonds Propres (RP)', 'FADEC Affecté (FA)',
                  'FADEC Non Affecté (FNA)', 'PTFs / Partenaires', 'Autres Fonds']
        values = [rp, fa, fn, ap, af]
        colors = ['#1F4E79', '#2E86AB', '#3BB273', '#E9C46A', '#E76F51']

        pairs = [(l, v, c) for l, v, c in zip(labels, values, colors) if v > 0]
        if not pairs:
            return None
        f_labels, f_values, f_colors = zip(*pairs)

        fig, ax = plt.subplots(figsize=(7, 4.5), dpi=110)
        wedges, _, autotexts = ax.pie(
            f_values, labels=None, colors=f_colors,
            autopct='%1.1f%%', startangle=90,
            wedgeprops=dict(width=0.55, edgecolor='white', linewidth=1.5)
        )
        for at in autotexts:
            at.set_fontsize(9)
        ax.legend(wedges, f_labels, loc='lower center',
                  bbox_to_anchor=(0.5, -0.28), ncol=2, fontsize=8.5, frameon=False)
        ax.set_title(title, fontsize=11, fontweight='bold', pad=14)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight', dpi=130)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _chart_directions(dirs_data, title='Activités et tâches par direction'):
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import numpy as np

        if not dirs_data:
            return None

        labels = [d['code'] for d in dirs_data]
        n_act  = [d['nb_activites'] for d in dirs_data]
        n_tch  = [d['nb_taches']    for d in dirs_data]

        y = np.arange(len(labels))
        h = 0.38

        fig_h = max(3.5, len(labels) * 0.45 + 1.2)
        fig, ax = plt.subplots(figsize=(9, fig_h), dpi=110)

        ax.barh(y + h / 2, n_act, h, label='Activités', color='#1F4E79', alpha=0.85)
        ax.barh(y - h / 2, n_tch, h, label='Tâches',    color='#3BB273', alpha=0.85)

        ax.set_yticks(y)
        ax.set_yticklabels(labels, fontsize=9)
        ax.set_xlabel('Nombre', fontsize=9)
        ax.set_title(title, fontsize=11, fontweight='bold')
        ax.legend(fontsize=9)
        ax.xaxis.set_major_locator(plt.MaxNLocator(integer=True))
        ax.invert_yaxis()
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight', dpi=130)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


# ── Construction du document Word ─────────────────────────────────────────────

def _build_word(annee, s, static_img_path, cibles=None, cibles_directions=None, cibles_services=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    cibles            = cibles or {}
    cibles_directions = cibles_directions or []
    cibles_services   = cibles_services   or []

    # ── Helpers XML / style ───────────────────────────────────────────────────

    def repeat_hdr(row):
        """Répéter la ligne d'en-tête sur chaque page pour un tableau."""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        hdr = OxmlElement('w:tblHeader')
        hdr.set(qn('w:val'), 'true')
        trPr.append(hdr)

    def shade(cell, hex_col):
        """Colorier le fond d'une cellule."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_col)
        tcPr.append(shd)

    def borders(table):
        """Ajouter des bordures fines à tous les côtés d'un tableau."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBrd = OxmlElement('w:tblBorders')
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'AAAAAA')
            tblBrd.append(b)
        tblPr.append(tblBrd)

    def set_col_width(table, col_idx, width_cm):
        """Fixer la largeur d'une colonne (en Cm)."""
        from docx.oxml import OxmlElement as OE
        from docx.shared import Cm as _Cm
        twips = int(width_cm * 567)  # 1 cm = 567 twips
        for row in table.rows:
            cell = row.cells[col_idx]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OE('w:tcW')
            tcW.set(qn('w:w'), str(twips))
            tcW.set(qn('w:type'), 'dxa')
            old = tcPr.find(qn('w:tcW'))
            if old is not None:
                tcPr.remove(old)
            tcPr.insert(0, tcW)

    # ── Fonctions de mise en page ─────────────────────────────────────────────

    def switch_to_landscape(doc):
        """Passer en orientation paysage (A4 : 29,7 × 21 cm, marges 1,5 cm)."""
        sec = doc.add_section()
        sec.orientation    = WD_ORIENT.LANDSCAPE
        sec.page_width     = Cm(29.7)
        sec.page_height    = Cm(21)
        sec.left_margin    = sec.right_margin  = Cm(1.5)
        sec.top_margin     = sec.bottom_margin = Cm(1.5)
        return sec

    def switch_to_portrait(doc):
        """Repasser en orientation portrait (A4 : 21 × 29,7 cm, marges standards)."""
        sec = doc.add_section()
        sec.orientation    = WD_ORIENT.PORTRAIT
        sec.page_width     = Cm(21)
        sec.page_height    = Cm(29.7)
        sec.left_margin    = sec.right_margin  = Cm(2.5)
        sec.top_margin     = Cm(2)
        sec.bottom_margin  = Cm(2)
        return sec

    # ── Fonctions de style de texte ───────────────────────────────────────────

    def h1(doc, text):
        """Titre de section principale (niveau 1)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        return p

    def h2(doc, text, color=(0x2E, 0x86, 0xAB)):
        """Titre de sous-section (niveau 2)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*color)
        return p

    def interp(doc, text):
        """Paragraphe d'interprétation / commentaire (italique, indentation)."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run('ℹ️  ' + text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        return p

    # ── Fonctions de construction de cellules ─────────────────────────────────

    def hdr_cell(cell, text, bg='1F4E79'):
        """Cellule d'en-tête de tableau (fond coloré, texte blanc, gras)."""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, bg)

    def data_cell(cell, text, align='left', bold=False, size=9, fg=None):
        """Cellule de données standard."""
        p = cell.paragraphs[0]
        p.alignment = {'left':   WD_ALIGN_PARAGRAPH.LEFT,
                       'center': WD_ALIGN_PARAGRAPH.CENTER,
                       'right':  WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(str(text) if text is not None else '')
        r.bold = bold
        r.font.size = Pt(size)
        if fg:
            r.font.color.rgb = RGBColor(*fg)

    def cible_cell(cell, val):
        """Cellule affichant un taux de cible trimestriel avec couleur selon niveau."""
        v = float(val) if val is not None else 0.0
        if   v < 25:  col = 'FADBD8'   # rouge clair  — moins du quart
        elif v < 50:  col = 'FDEBD0'   # orange clair — moins de la moitié
        elif v < 75:  col = 'D6EAF8'   # bleu clair   — moins des trois quarts
        else:         col = 'D5F5E3'   # vert clair   — objectif atteint ou proche
        data_cell(cell, f"{v:.1f} %", align='center', bold=True, size=9)
        shade(cell, col)

    # ── Fonctions de construction de tableaux ─────────────────────────────────

    def table_prog(doc, data, nb_act, nb_tch, tot_b, row_col='FFB6C1'):
        """Tableau synthétique programme → projet (6 colonnes)."""
        n_rows = 1 + sum(1 + len(p['projets']) for p in data) + 1
        tbl = doc.add_table(rows=max(n_rows, 2), cols=6)
        borders(tbl)
        for i, txt in enumerate(['Code', 'Intitulé', 'Activités', 'Tâches', 'Poids (%)', 'Budget (F CFA)']):
            hdr_cell(tbl.cell(0, i), txt, bg='1F4E79')
        repeat_hdr(tbl.rows[0])
        ri = 1
        for prog in data:
            r = tbl.rows[ri]
            for ci, (val, al, bd) in enumerate([
                (prog['code'],                               'center', True),
                (f"Programme {prog['code']} : {prog['nom']}", 'left',  True),
                (str(prog['nb_activites']),                  'center', True),
                (str(prog['nb_taches']),                     'center', True),
                (f"{prog['poids']:.1f}",                     'center', True),
                (_fmt(prog['budget']),                        'right',  True),
            ]):
                data_cell(r.cells[ci], val, align=al, bold=bd, size=9)
                shade(r.cells[ci], 'FFFF99')
            ri += 1
            for proj in prog['projets']:
                r = tbl.rows[ri]
                for ci, (val, al, bd) in enumerate([
                    (proj['code'],                               'center', False),
                    (f"  Projet {proj['code']} : {proj['nom']}", 'left',  False),
                    (str(proj['nb_activites']),                  'center', False),
                    (str(proj['nb_taches']),                     'center', False),
                    (f"{proj['poids']:.1f}",                     'center', False),
                    (_fmt(proj['budget']),                        'right',  False),
                ]):
                    data_cell(r.cells[ci], val, align=al, bold=bd, size=9)
                    shade(r.cells[ci], row_col)
                ri += 1
        # Ligne total
        r = tbl.rows[ri]
        for ci, (val, al, bd) in enumerate([
            ('TOTAL', 'center', True), ('', 'left', True),
            (str(nb_act), 'center', True), (str(nb_tch), 'center', True),
            ('', 'center', True), (_fmt(tot_b), 'right', True),
        ]):
            data_cell(r.cells[ci], val, align=al, bold=bd, size=9)
            shade(r.cells[ci], 'AED6F1')
        return tbl

    def table_sources(doc, rp_fa_fn_ap_af_tot, pct_fn, bg_cols):
        """Tableau répartition budgétaire par source de financement (3 colonnes)."""
        SRC_LABELS = ['Fonds Propres (RP)', 'FADEC Affecté (FA)',
                      'FADEC Non Affecté (FNA)', 'PTFs / Partenaires', 'Autres Fonds']
        rp, fa, fn, ap, af, total = rp_fa_fn_ap_af_tot
        tbl = doc.add_table(rows=len(SRC_LABELS) + 2, cols=3)
        borders(tbl)
        for i, txt in enumerate(['Source de financement', 'Montant (F CFA)', 'Part (%)']):
            hdr_cell(tbl.cell(0, i), txt, bg='2C3E50')
        repeat_hdr(tbl.rows[0])
        for i, (lbl, val, col) in enumerate(zip(SRC_LABELS, [rp, fa, fn, ap, af], bg_cols), 1):
            r = tbl.rows[i]
            data_cell(r.cells[0], lbl, size=9)
            data_cell(r.cells[1], _fmt(val), align='right', size=9)
            data_cell(r.cells[2], f"{pct_fn(val)} %", align='center', size=9)
            for cell in r.cells:
                shade(cell, col)
        tot_row = tbl.rows[-1]
        data_cell(tot_row.cells[0], 'TOTAL', bold=True, size=9)
        data_cell(tot_row.cells[1], _fmt(total), align='right', bold=True, size=9)
        data_cell(tot_row.cells[2], '100 %', align='center', bold=True, size=9)
        for cell in tot_row.cells:
            shade(cell, 'AED6F1')
        return tbl

    def table_directions(doc, dirs_data, nb_act_total, nb_tch_total):
        """Tableau activités/tâches par direction responsable (4 colonnes)."""
        tbl = doc.add_table(rows=len(dirs_data) + 2, cols=4)
        borders(tbl)
        for i, txt in enumerate(['Code', 'Direction', 'Nb activités (resp.)', 'Nb tâches (resp.)']):
            hdr_cell(tbl.cell(0, i), txt, bg='27AE60' if i >= 2 else '2C3E50')
        repeat_hdr(tbl.rows[0])
        for i, d in enumerate(dirs_data, 1):
            r = tbl.rows[i]
            data_cell(r.cells[0], d['code'], align='center', size=9)
            data_cell(r.cells[1], d['nom'],  align='left',   size=9)
            data_cell(r.cells[2], str(d['nb_activites']), align='center',
                      bold=(d['nb_activites'] > 0), size=9)
            data_cell(r.cells[3], str(d['nb_taches']),    align='center',
                      bold=(d['nb_taches'] > 0),    size=9)
            col = 'F9F9F9' if (d['nb_activites'] == 0 and d['nb_taches'] == 0) else 'EBF5FB'
            for cell in r.cells:
                shade(cell, col)
        tot = tbl.rows[-1]
        for ci, (val, al, bd) in enumerate([
            ('—', 'center', False), ('TOTAL', 'left', True),
            (str(nb_act_total), 'center', True), (str(nb_tch_total), 'center', True),
        ]):
            data_cell(tot.cells[ci], val, align=al, bold=bd, size=9)
            shade(tot.cells[ci], 'AED6F1')

    def table_activites_detail(doc, liste_act, tot_src, hdr_bg='C0392B',
                               pair_col='FDEBD0', impair_col='FEF9E7'):
        """Tableau détaillé des activités en paysage (11 colonnes + totaux).
        Utilisé pour la liste investissement et fonctionnement.
        """
        cols = ['Code', 'Désignation', 'Dir.\nresp.', 'Période', 'Poids\n(%)',
                'RP', 'FADEC\nAff.', 'FADEC\nN-Aff.', 'PTFs', 'Autres', 'Total (F CFA)']
        n_rows = len(liste_act) + 2
        tbl = doc.add_table(rows=n_rows, cols=len(cols))
        borders(tbl)
        for i, txt in enumerate(cols):
            hdr_cell(tbl.cell(0, i), txt, bg=hdr_bg)
        repeat_hdr(tbl.rows[0])

        rp_tot = fa_tot = fn_tot = ap_tot = af_tot = bud_tot = 0
        for ri, a in enumerate(liste_act, 1):
            r = tbl.rows[ri]
            col_bg = pair_col if ri % 2 == 0 else impair_col
            for ci, (val, al, bd) in enumerate([
                (a['code'],            'center', False),
                (a['nom'],             'left',   False),
                (a['direction'],       'center', False),
                (a['periode'],         'center', False),
                (f"{a['poids']:.1f}",  'center', False),
                (_fmt(a['rp']),        'right',  False),
                (_fmt(a['fa']),        'right',  False),
                (_fmt(a['fn']),        'right',  False),
                (_fmt(a['ap']),        'right',  False),
                (_fmt(a['af']),        'right',  False),
                (_fmt(a['budget']),    'right',  True),
            ]):
                data_cell(r.cells[ci], val, align=al, bold=bd, size=8)
                shade(r.cells[ci], col_bg)
            rp_tot  += a['rp']; fa_tot += a['fa']; fn_tot += a['fn']
            ap_tot  += a['ap']; af_tot += a['af']; bud_tot += a['budget']

        # Ligne de totaux
        tot_r = tbl.rows[-1]
        for ci, (val, al) in enumerate([
            ('TOTAL', 'center'), ('', 'left'), ('', 'center'), ('', 'center'), ('', 'center'),
            (_fmt(tot_rp := tot_src[0]),  'right'),
            (_fmt(tot_src[1]), 'right'),
            (_fmt(tot_src[2]), 'right'),
            (_fmt(tot_src[3]), 'right'),
            (_fmt(tot_src[4]), 'right'),
            (_fmt(tot_src[5]), 'right'),
        ]):
            data_cell(tot_r.cells[ci], val, align=al, bold=True, size=8)
            shade(tot_r.cells[ci], 'AED6F1')
        return tbl

    def table_cibles_prog(doc, cibles):
        """Tableau des cibles trimestrielles par programme et projet (5 colonnes)."""
        programmes_c = cibles.get('programmes', [])
        global_c     = cibles.get('global', {1: 0, 2: 0, 3: 0, 4: 0})
        # Compter les lignes : 1 (global) + 1 par prog + 1 par proj + note
        n = 1 + sum(1 + len(p.get('projets', [])) for p in programmes_c)
        tbl = doc.add_table(rows=n + 1, cols=5)
        borders(tbl)
        for i, txt in enumerate(['Niveau', 'T1 (fin Mars)', 'T2 (fin Juin)', 'T3 (fin Sept.)', 'T4 (fin Déc.)']):
            hdr_cell(tbl.cell(0, i), txt, bg='1A5276')
        repeat_hdr(tbl.rows[0])
        ri = 1
        # Ligne PTA Global
        r = tbl.rows[ri]
        data_cell(r.cells[0], 'PTA Global', bold=True, size=9)
        for ti, trim in enumerate((1, 2, 3, 4), 1):
            cible_cell(r.cells[ti], global_c.get(trim, 0))
        shade(r.cells[0], 'FFFF99')
        ri += 1
        # Programmes et leurs projets
        for prog in programmes_c:
            r = tbl.rows[ri]
            data_cell(r.cells[0], f"Programme {prog['code']} — {prog['nom']}", bold=True, size=9)
            for ti, trim in enumerate((1, 2, 3, 4), 1):
                cible_cell(r.cells[ti], prog['cibles'].get(trim, 0))
            shade(r.cells[0], 'D4E6F1')
            ri += 1
            for proj in prog.get('projets', []):
                r = tbl.rows[ri]
                data_cell(r.cells[0], f"   {proj['code']} — {proj['nom']}", size=9)
                for ti, trim in enumerate((1, 2, 3, 4), 1):
                    cible_cell(r.cells[ti], proj['cibles'].get(trim, 0))
                shade(r.cells[0], 'EBF5FB')
                ri += 1
        return tbl

    def table_cibles_dirs(doc, cibles_dirs):
        """Tableau des cibles trimestrielles par direction (7 colonnes).
        Colonnes : Code | Direction | Nb tâches | T1 | T2 | T3 | T4
        """
        tbl = doc.add_table(rows=len(cibles_dirs) + 1, cols=7)
        borders(tbl)
        for i, txt in enumerate(['Code', 'Direction', 'Nb tâches',
                                  'T1 (fin Mars)', 'T2 (fin Juin)',
                                  'T3 (fin Sept.)', 'T4 (fin Déc.)']):
            hdr_cell(tbl.cell(0, i), txt, bg='1A5276')
        repeat_hdr(tbl.rows[0])
        for ri, d in enumerate(cibles_dirs, 1):
            r = tbl.rows[ri]
            data_cell(r.cells[0], d['code'],           align='center', bold=True, size=9)
            data_cell(r.cells[1], d['nom'],             align='left',             size=9)
            data_cell(r.cells[2], str(d['nb_taches']), align='center',            size=9)
            # Colonnes T1→T4 : indices 3,4,5,6 ; clés du dict cibles : 1,2,3,4
            for trim in (1, 2, 3, 4):
                cible_cell(r.cells[trim + 2], d['cibles'].get(trim, 0))
            col = 'F4F6F7' if ri % 2 == 0 else 'EBF5FB'
            for ci in range(3):
                shade(r.cells[ci], col)
        return tbl

    def table_cibles_svcs(doc, cibles_svcs):
        """Tableau des cibles trimestrielles par service (8 colonnes).
        Colonnes : Dir. | Code | Service | Nb tâches | T1 | T2 | T3 | T4
        """
        tbl = doc.add_table(rows=len(cibles_svcs) + 1, cols=8)
        borders(tbl)
        for i, txt in enumerate(['Dir.', 'Code', 'Service', 'Nb tâches',
                                  'T1 (fin Mars)', 'T2 (fin Juin)',
                                  'T3 (fin Sept.)', 'T4 (fin Déc.)']):
            hdr_cell(tbl.cell(0, i), txt, bg='1A5276')
        repeat_hdr(tbl.rows[0])
        for ri, sv in enumerate(cibles_svcs, 1):
            r = tbl.rows[ri]
            data_cell(r.cells[0], sv['direction_code'], align='center', bold=True, size=9)
            data_cell(r.cells[1], sv['code'],           align='center', bold=True, size=9)
            data_cell(r.cells[2], sv['nom'],             align='left',             size=9)
            data_cell(r.cells[3], str(sv['nb_taches']), align='center',            size=9)
            # Colonnes T1→T4 : indices 4,5,6,7 ; clés du dict cibles : 1,2,3,4
            for trim in (1, 2, 3, 4):
                cible_cell(r.cells[trim + 3], sv['cibles'].get(trim, 0))
            col = 'F4F6F7' if ri % 2 == 0 else 'EBF5FB'
            for ci in range(4):
                shade(r.cells[ci], col)
        return tbl

    # ── Helpers pour commentaires dynamiques ─────────────────────────────────

    def _src_principale(rp, fa, fn, ap, af):
        """Retourne (nom_affichable, valeur) de la source la plus importante."""
        candidats = [
            ('les Ressources Propres (RP)', rp),
            ('le FADEC Affecté (FA)', fa),
            ('le FADEC Non Affecté (FNA)', fn),
            ('les PTFs/Partenaires', ap),
            ('les Autres Fonds', af),
        ]
        actifs = [(n, v) for n, v in candidats if v > 0]
        return max(actifs, key=lambda x: x[1]) if actifs else ('—', 0)

    def _commentaire_sources(rp, fa, fn, ap, af, total, pct_fn):
        """Génère un commentaire dynamique sur la répartition des sources."""
        if total == 0:
            return "Aucun budget renseigné pour cette catégorie."
        nom_src, val_src = _src_principale(rp, fa, fn, ap, af)
        pct_src  = pct_fn(val_src)
        ext_val  = total - rp
        pct_ext  = round(ext_val / total * 100, 1)
        pct_rp   = round(rp / total * 100, 1)
        parties  = []
        parties.append(
            f"La principale source de financement est {nom_src} "
            f"avec {_fmt(val_src)} F CFA ({pct_src} % du budget)."
        )
        if pct_rp >= 50:
            parties.append(
                f"La commune finance {pct_rp} % de ce budget sur ses ressources propres."
            )
        elif pct_ext >= 50:
            parties.append(
                f"Les financements extérieurs (PTFs, FADEC, partenaires) couvrent "
                f"{pct_ext} % du budget, témoignant d'un fort appui externe."
            )
        return ' '.join(parties)

    # ── Construction du document ──────────────────────────────────────────────

    doc = Document()

    # Section 0 : portrait A4, sans en-tête ni pied de page
    sec0 = doc.sections[0]
    sec0.orientation    = WD_ORIENT.PORTRAIT
    sec0.page_width     = Cm(21)
    sec0.page_height    = Cm(29.7)
    sec0.left_margin    = sec0.right_margin  = Cm(2.5)
    sec0.top_margin     = Cm(2.5)
    sec0.bottom_margin  = Cm(2)

    # ── Page de titre ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('RAPPORT STATISTIQUE')
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'PLAN DE TRAVAIL ANNUEL — Exercice {annee.annee}')
    r2.bold = True; r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Mairie d'Adja-Ouèrè")
    r3.bold = True; r3.font.size = Pt(12)

    logo_path = os.path.join(static_img_path, 'logo_commune.png')
    if os.path.exists(logo_path):
        pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.add_run().add_picture(logo_path, height=Cm(2.4))

    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')
    r4 = p4.add_run(f'Généré le {date_str}')
    r4.italic = True; r4.font.size = Pt(9)
    r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.add_run('─' * 60).font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    SRC_COLS = ['EBF5FB', 'D6EAF8', 'D5F5E3', 'FEFDE7', 'FDEDEC']

    # ════════════════════════════════════════════════════════════════════════════
    # I. SYNTHÈSE GLOBALE
    # ════════════════════════════════════════════════════════════════════════════
    h1(doc, 'I. SYNTHÈSE GLOBALE')

    # Tableau indicateurs clés
    tbl_kpi = doc.add_table(rows=2, cols=5)
    borders(tbl_kpi)
    HDRS = ['Budget total PTA', 'Programmes', 'Projets', 'Activités', 'Tâches']
    VALS = [f"{_fmt(s['total_budget'])} F CFA",
            str(s['nb_programmes']), str(s['nb_projets']),
            str(s['nb_activites']),  str(s['nb_taches'])]
    for i, (h_txt, v_txt) in enumerate(zip(HDRS, VALS)):
        hdr_cell(tbl_kpi.cell(0, i), h_txt)
        c = tbl_kpi.cell(1, i)
        data_cell(c, v_txt, align='center', bold=True, size=11)
        shade(c, 'D6EAF8')
    repeat_hdr(tbl_kpi.rows[0])
    doc.add_paragraph()

    # Tableau répartition par nature d'activité
    h2(doc, "Répartition par nature d'activité")
    tbl_nat = doc.add_table(rows=4, cols=4)
    borders(tbl_nat)
    for i, txt in enumerate(['Nature', 'Activités', 'Tâches', 'Budget (F CFA)']):
        hdr_cell(tbl_nat.cell(0, i), txt, bg='2C3E50')
    repeat_hdr(tbl_nat.rows[0])
    for ri, (lbl, na, nt, bud, col) in enumerate([
        ("Toutes activités",            s['nb_activites'],        s['nb_taches'],        s['total_budget'], 'D6EAF8'),
        ("Activités d'investissement",  s['nb_activites_invest'], s['nb_taches_invest'], s['inv_budget'],   'FDEBD0'),
        ("Activités de fonctionnement", s['nb_activites_fonct'],  s['nb_taches_fonct'],  s['fct_budget'],   'D5F5E3'),
    ], 1):
        data_cell(tbl_nat.rows[ri].cells[0], lbl,       bold=True, size=9)
        data_cell(tbl_nat.rows[ri].cells[1], str(na),   align='center', bold=True, size=9)
        data_cell(tbl_nat.rows[ri].cells[2], str(nt),   align='center', bold=True, size=9)
        data_cell(tbl_nat.rows[ri].cells[3], _fmt(bud), align='right',  bold=True, size=9)
        for cell in tbl_nat.rows[ri].cells:
            shade(cell, col)
    doc.add_paragraph()

    # ── Commentaire dynamique I ───────────────────────────────────────────────
    pct_inv = round(s['inv_budget'] / s['total_budget'] * 100, 1) if s['total_budget'] else 0
    pct_fct = round(s['fct_budget'] / s['total_budget'] * 100, 1) if s['total_budget'] else 0
    avg_bud = round(s['total_budget'] / s['nb_activites']) if s['nb_activites'] else 0
    avg_tch = round(s['nb_taches'] / s['nb_activites'], 1) if s['nb_activites'] else 0
    interp(doc,
        f"Le PTA {annee.annee} comprend {s['nb_programmes']} programme(s) répartis en "
        f"{s['nb_projets']} projet(s), {s['nb_activites']} activité(s) et "
        f"{s['nb_taches']} tâche(s). Le budget global s'élève à "
        f"{_fmt(s['total_budget'])} F CFA, dont {pct_inv} % consacré à l'investissement "
        f"({_fmt(s['inv_budget'])} F CFA) et {pct_fct} % au fonctionnement "
        f"({_fmt(s['fct_budget'])} F CFA). "
        f"Chaque activité dispose en moyenne d'un budget de {_fmt(avg_bud)} F CFA "
        f"et comprend en moyenne {avg_tch} tâche(s)."
    )

    # ════════════════════════════════════════════════════════════════════════════
    # II. RÉPARTITION BUDGÉTAIRE PAR SOURCE DE FINANCEMENT
    # ════════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    h1(doc, 'II. RÉPARTITION BUDGÉTAIRE PAR SOURCE DE FINANCEMENT')

    h2(doc, 'A. Toutes activités')
    table_sources(doc,
        (s['total_rp'], s['total_fa'], s['total_fn'], s['total_ap'], s['total_af'], s['total_budget']),
        s['pct'], SRC_COLS)
    chart_src = _chart_sources(s['total_rp'], s['total_fa'], s['total_fn'], s['total_ap'], s['total_af'])
    if chart_src:
        doc.add_picture(chart_src, width=Cm(12))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    # ── Commentaire dynamique II-A ────────────────────────────────────────────
    interp(doc, _commentaire_sources(
        s['total_rp'], s['total_fa'], s['total_fn'], s['total_ap'], s['total_af'],
        s['total_budget'], s['pct']
    ))

    if s['inv_budget'] > 0:
        h2(doc, "B. Activités d'investissement", color=(0xC0, 0x39, 0x2B))
        table_sources(doc,
            (s['inv_rp'], s['inv_fa'], s['inv_fn'], s['inv_ap'], s['inv_af'], s['inv_budget']),
            s['pct_invest'], SRC_COLS)
        chart_inv = _chart_sources(s['inv_rp'], s['inv_fa'], s['inv_fn'], s['inv_ap'], s['inv_af'],
                                   "Sources — Investissement")
        if chart_inv:
            doc.add_picture(chart_inv, width=Cm(11))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        # ── Commentaire dynamique II-B ────────────────────────────────────────
        interp(doc, _commentaire_sources(
            s['inv_rp'], s['inv_fa'], s['inv_fn'], s['inv_ap'], s['inv_af'],
            s['inv_budget'], s['pct_invest']
        ))

    if s['fct_budget'] > 0:
        h2(doc, 'C. Activités de fonctionnement', color=(0x1A, 0x72, 0x45))
        table_sources(doc,
            (s['fct_rp'], s['fct_fa'], s['fct_fn'], s['fct_ap'], s['fct_af'], s['fct_budget']),
            s['pct_fonct'], SRC_COLS)
        chart_fct = _chart_sources(s['fct_rp'], s['fct_fa'], s['fct_fn'], s['fct_ap'], s['fct_af'],
                                   "Sources — Fonctionnement")
        if chart_fct:
            doc.add_picture(chart_fct, width=Cm(11))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        # ── Commentaire dynamique II-C ────────────────────────────────────────
        interp(doc, _commentaire_sources(
            s['fct_rp'], s['fct_fa'], s['fct_fn'], s['fct_ap'], s['fct_af'],
            s['fct_budget'], s['pct_fonct']
        ))

    # ════════════════════════════════════════════════════════════════════════════
    # III. ANALYSE PAR PROGRAMME ET PAR PROJET
    # ════════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    h1(doc, 'III. ANALYSE PAR PROGRAMME ET PAR PROJET')

    h2(doc, 'A. Toutes activités')
    table_prog(doc, s['stats_par_programme'],
               s['nb_activites'], s['nb_taches'], s['total_budget'])
    doc.add_paragraph()
    # ── Commentaire dynamique III-A ───────────────────────────────────────────
    progs_all = [p for p in s['stats_par_programme'] if p['nb_activites'] > 0]
    if progs_all and s['total_budget'] > 0:
        top_bud = max(progs_all, key=lambda p: p['budget'])
        top_act = max(progs_all, key=lambda p: p['nb_activites'])
        top_tch = max(progs_all, key=lambda p: p['nb_taches'])
        pct_top = round(top_bud['budget'] / s['total_budget'] * 100, 1)
        _c = [
            f"Le Programme {top_bud['code']} ({top_bud['nom']}) concentre "
            f"la part la plus importante du budget avec {_fmt(top_bud['budget'])} F CFA "
            f"soit {pct_top} % du total."
        ]
        if top_act['code'] != top_bud['code']:
            _c.append(
                f"C'est le Programme {top_act['code']} qui compte le plus d'activités "
                f"({top_act['nb_activites']})."
            )
        if top_tch['code'] != top_act['code']:
            _c.append(
                f"Le Programme {top_tch['code']} regroupe le plus grand nombre de tâches "
                f"({top_tch['nb_taches']})."
            )
        if len(progs_all) == 1:
            _c.append("L'intégralité du budget est portée par un seul programme.")
        interp(doc, ' '.join(_c))

    if s['stats_par_programme_invest']:
        h2(doc, "B. Activités d'investissement", color=(0xC0, 0x39, 0x2B))
        table_prog(doc, s['stats_par_programme_invest'],
                   s['nb_activites_invest'], s['nb_taches_invest'],
                   s['inv_budget'], row_col='FDEBD0')
        doc.add_paragraph()
        # ── Commentaire dynamique III-B ───────────────────────────────────────
        progs_inv = s['stats_par_programme_invest']
        if progs_inv and s['inv_budget'] > 0:
            top_i = max(progs_inv, key=lambda p: p['budget'])
            pct_i = round(top_i['budget'] / s['inv_budget'] * 100, 1)
            interp(doc,
                f"En matière d'investissement, le Programme {top_i['code']} ({top_i['nom']}) "
                f"est le plus chargé avec {_fmt(top_i['budget'])} F CFA ({pct_i} % du budget "
                f"d'investissement) et {top_i['nb_activites']} activité(s)."
            )

    if s['stats_par_programme_fonct']:
        h2(doc, 'C. Activités de fonctionnement', color=(0x1A, 0x72, 0x45))
        table_prog(doc, s['stats_par_programme_fonct'],
                   s['nb_activites_fonct'], s['nb_taches_fonct'],
                   s['fct_budget'], row_col='D5F5E3')
        doc.add_paragraph()
        # ── Commentaire dynamique III-C ───────────────────────────────────────
        progs_fct = s['stats_par_programme_fonct']
        if progs_fct and s['fct_budget'] > 0:
            top_f = max(progs_fct, key=lambda p: p['nb_activites'])
            interp(doc,
                f"Pour le fonctionnement, le Programme {top_f['code']} ({top_f['nom']}) "
                f"regroupe le plus grand nombre d'activités ({top_f['nb_activites']}) "
                f"pour un budget de {_fmt(top_f['budget'])} F CFA."
            )

    # ════════════════════════════════════════════════════════════════════════════
    # IV. ACTIVITÉS ET TÂCHES PAR DIRECTION
    # ════════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    h1(doc, 'IV. ACTIVITÉS ET TÂCHES PAR DIRECTION RESPONSABLE')

    h2(doc, 'A. Toutes activités')
    table_directions(doc, s['stats_directions'], s['nb_activites'], s['nb_taches'])
    chart_dir = _chart_directions(s['stats_directions'])
    if chart_dir:
        doc.add_picture(chart_dir, width=Cm(14))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    # ── Commentaire dynamique IV-A ────────────────────────────────────────────
    dirs_actives = [d for d in s['stats_directions'] if d['nb_activites'] > 0]
    dirs_inactives = [d for d in s['stats_directions'] if d['nb_activites'] == 0]
    if dirs_actives:
        top_da = max(dirs_actives, key=lambda d: d['nb_activites'])
        _c = [
            f"{len(dirs_actives)} direction(s) sur {len(s['stats_directions'])} "
            f"sont responsables d'au moins une activité dans le PTA {annee.annee}."
        ]
        _c.append(
            f"La direction la plus sollicitée est {top_da['code']} avec "
            f"{top_da['nb_activites']} activité(s) et {top_da['nb_taches']} tâche(s)."
        )
        if dirs_inactives:
            codes_inact = ', '.join(d['code'] for d in dirs_inactives[:6])
            suite = '…' if len(dirs_inactives) > 6 else ''
            _c.append(
                f"{len(dirs_inactives)} direction(s) n'ont aucune activité assignée : "
                f"{codes_inact}{suite}."
            )
        interp(doc, ' '.join(_c))

    if s['nb_activites_invest'] > 0:
        h2(doc, "B. Activités d'investissement", color=(0xC0, 0x39, 0x2B))
        table_directions(doc, s['stats_directions_invest'],
                         s['nb_activites_invest'], s['nb_taches_invest'])
        chart_dir_inv = _chart_directions(s['stats_directions_invest'],
                                          "Par direction — Investissement")
        if chart_dir_inv:
            doc.add_picture(chart_dir_inv, width=Cm(13))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        # ── Commentaire dynamique IV-B ────────────────────────────────────────
        dirs_inv_act = [d for d in s['stats_directions_invest'] if d['nb_activites'] > 0]
        if dirs_inv_act:
            top_di = max(dirs_inv_act, key=lambda d: d['nb_activites'])
            interp(doc,
                f"{len(dirs_inv_act)} direction(s) pilotent des activités d'investissement. "
                f"La plus active est {top_di['code']} avec {top_di['nb_activites']} "
                f"activité(s) d'investissement et {top_di['nb_taches']} tâche(s)."
            )

    if s['nb_activites_fonct'] > 0:
        h2(doc, 'C. Activités de fonctionnement', color=(0x1A, 0x72, 0x45))
        table_directions(doc, s['stats_directions_fonct'],
                         s['nb_activites_fonct'], s['nb_taches_fonct'])
        chart_dir_fct = _chart_directions(s['stats_directions_fonct'],
                                          "Par direction — Fonctionnement")
        if chart_dir_fct:
            doc.add_picture(chart_dir_fct, width=Cm(13))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        # ── Commentaire dynamique IV-C ────────────────────────────────────────
        dirs_fct_act = [d for d in s['stats_directions_fonct'] if d['nb_activites'] > 0]
        if dirs_fct_act:
            top_df = max(dirs_fct_act, key=lambda d: d['nb_activites'])
            interp(doc,
                f"{len(dirs_fct_act)} direction(s) pilotent des activités de fonctionnement. "
                f"La plus chargée est {top_df['code']} avec {top_df['nb_activites']} "
                f"activité(s) de fonctionnement et {top_df['nb_taches']} tâche(s)."
            )

    # ════════════════════════════════════════════════════════════════════════════
    # V. LISTE DES ACTIVITÉS D'INVESTISSEMENT  [paysage]
    # ════════════════════════════════════════════════════════════════════════════
    if s.get('activites_investissement'):
        ai_list = s['activites_investissement']
        switch_to_landscape(doc)
        h1(doc, "V. LISTE DES ACTIVITÉS D'INVESTISSEMENT")
        # ── Commentaire dynamique V ───────────────────────────────────────────
        top_ai  = max(ai_list, key=lambda a: a['budget'])
        avg_ai  = round(s['inv_budget'] / len(ai_list)) if ai_list else 0
        dirs_ai = len(set(a['direction'] for a in ai_list))
        pct_ai  = round(top_ai['budget'] / s['inv_budget'] * 100, 1) if s['inv_budget'] else 0
        interp(doc,
            f"Les {len(ai_list)} activité(s) d'investissement relèvent de {dirs_ai} direction(s) "
            f"pour un budget total de {_fmt(s['inv_budget'])} F CFA "
            f"(budget moyen par activité : {_fmt(avg_ai)} F CFA). "
            f"L'activité la plus importante est « {top_ai['nom']} » ({top_ai['code']}, "
            f"{top_ai['direction']}) avec {_fmt(top_ai['budget'])} F CFA, "
            f"soit {pct_ai} % du total investissement. "
            f"Légende colonnes : RP = Ressources Propres, FA = FADEC Affecté, "
            f"FNA = FADEC Non Affecté, AP = PTFs/Partenaires, AF = Autres Fonds."
        )
        table_activites_detail(
            doc, ai_list,
            (s['inv_rp'], s['inv_fa'], s['inv_fn'], s['inv_ap'], s['inv_af'], s['inv_budget']),
            hdr_bg='C0392B', pair_col='FDEBD0', impair_col='FEF9E7'
        )
        switch_to_portrait(doc)

    # ════════════════════════════════════════════════════════════════════════════
    # VI. CIBLES TRIMESTRIELLES D'EXÉCUTION
    # ════════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    h1(doc, "VI. CIBLES TRIMESTRIELLES D'EXÉCUTION")
    interp(doc,
        "Les cibles représentent le taux d'avancement attendu si chaque tâche est réalisée "
        "à 100 % dans ses délais planifiés. Calcul : pro-ratisation linéaire mensuelle sur "
        "la période [début → fin] de chaque tâche, remontée hiérarchique par poids renormalisés "
        "(tâche → activité → projet → programme). Tâches sans période planifiée : "
        "répartition uniforme 25 % / 50 % / 75 % / 100 %. "
        "Code couleur — Rouge : < 25 % | Orange : 25–50 % | Bleu : 50–75 % | Vert : ≥ 75 %."
    )
    doc.add_paragraph()

    # ── VI-A : Par Programme / Projet ─────────────────────────────────────────
    if cibles:
        h2(doc, 'A. Par Programme et par Projet')
        table_cibles_prog(doc, cibles)
        doc.add_paragraph()
        # ── Commentaire dynamique VI-A ────────────────────────────────────────
        g = cibles.get('global', {1: 0, 2: 0, 3: 0, 4: 0})
        progs_c = cibles.get('programmes', [])
        # T4 = 100 % par construction (toutes tâches réalisées dans leurs délais).
        # L'intérêt est la courbe trimestrielle : T1, T2, T3 révèlent la saisonnalité.
        _c = [
            f"Selon la planification actuelle, le PTA {annee.annee} devrait atteindre "
            f"{g.get(1, 0):.1f} % d'exécution à fin mars (T1), "
            f"{g.get(2, 0):.1f} % à fin juin (T2) et "
            f"{g.get(3, 0):.1f} % à fin septembre (T3)."
        ]
        if len(progs_c) > 1:
            top_p1 = max(progs_c, key=lambda p: p['cibles'].get(1, 0))
            bot_p1 = min(progs_c, key=lambda p: p['cibles'].get(1, 0))
            top_p3 = max(progs_c, key=lambda p: p['cibles'].get(3, 0))
            bot_p3 = min(progs_c, key=lambda p: p['cibles'].get(3, 0))
            _c.append(
                f"Dès le premier trimestre, le Programme {top_p1['code']} est le plus avancé "
                f"({top_p1['cibles'].get(1, 0):.1f} %) ; "
                f"le Programme {bot_p1['code']} est le moins avancé "
                f"({bot_p1['cibles'].get(1, 0):.1f} %)."
            )
            if top_p3['code'] != top_p1['code'] or bot_p3['code'] != bot_p1['code']:
                _c.append(
                    f"À fin T3, le Programme {top_p3['code']} mène avec "
                    f"{top_p3['cibles'].get(3, 0):.1f} % ; "
                    f"le Programme {bot_p3['code']} affiche {bot_p3['cibles'].get(3, 0):.1f} %."
                )
        interp(doc, ' '.join(_c))

    # ── VI-B : Par Direction ──────────────────────────────────────────────────
    if cibles_directions:
        h2(doc, 'B. Par Direction responsable')
        if len(cibles_directions) > 12:
            switch_to_landscape(doc)
        table_cibles_dirs(doc, cibles_directions)
        if len(cibles_directions) > 12:
            switch_to_portrait(doc)
        else:
            doc.add_paragraph()
        # ── Commentaire dynamique VI-B ────────────────────────────────────────
        # T4 = 100 % par construction. On compare T1 et T3 pour détecter
        # les directions à charge précoce vs tardive.
        avg_t1d = sum(d['cibles'].get(1, 0) for d in cibles_directions) / len(cibles_directions)
        avg_t3d = sum(d['cibles'].get(3, 0) for d in cibles_directions) / len(cibles_directions)
        top_d1  = max(cibles_directions, key=lambda d: d['cibles'].get(1, 0))
        bot_d1  = min(cibles_directions, key=lambda d: d['cibles'].get(1, 0))
        top_d3  = max(cibles_directions, key=lambda d: d['cibles'].get(3, 0))
        bot_d3  = min(cibles_directions, key=lambda d: d['cibles'].get(3, 0))
        _c = [
            f"{len(cibles_directions)} direction(s) ont des tâches planifiées dans ce PTA. "
            f"En moyenne, {avg_t1d:.1f} % des tâches devraient être exécutées à fin T1 "
            f"et {avg_t3d:.1f} % à fin T3."
        ]
        _c.append(
            f"La direction la plus avancée dès T1 est {top_d1['code']} "
            f"({top_d1['cibles'].get(1, 0):.1f} %), signe d'activités concentrées "
            f"en début d'année."
        )
        if len(cibles_directions) > 1 and bot_d1['code'] != top_d1['code']:
            _c.append(
                f"À l'inverse, {bot_d1['code']} n'atteint que "
                f"{bot_d1['cibles'].get(1, 0):.1f} % en T1, "
                f"avec des activités davantage concentrées en fin d'année."
            )
        if len(cibles_directions) > 1 and (top_d3['code'] != top_d1['code'] or bot_d3['code'] != bot_d1['code']):
            _c.append(
                f"À fin T3, {top_d3['code']} est la plus avancée ({top_d3['cibles'].get(3, 0):.1f} %) "
                f"et {bot_d3['code']} la moins avancée ({bot_d3['cibles'].get(3, 0):.1f} %)."
            )
        interp(doc, ' '.join(_c))

    # ── VI-C : Par Service ────────────────────────────────────────────────────
    if cibles_services:
        h2(doc, 'C. Par Service')
        if len(cibles_services) > 10:
            switch_to_landscape(doc)
        table_cibles_svcs(doc, cibles_services)
        if len(cibles_services) > 10:
            switch_to_portrait(doc)
        else:
            doc.add_paragraph()
        # ── Commentaire dynamique VI-C ────────────────────────────────────────
        # T4 = 100 % par construction. On analyse T1 et T3.
        avg_t1s = sum(sv['cibles'].get(1, 0) for sv in cibles_services) / len(cibles_services)
        avg_t3s = sum(sv['cibles'].get(3, 0) for sv in cibles_services) / len(cibles_services)
        top_s1  = max(cibles_services, key=lambda sv: sv['cibles'].get(1, 0))
        bot_s1  = min(cibles_services, key=lambda sv: sv['cibles'].get(1, 0))
        top_s3  = max(cibles_services, key=lambda sv: sv['cibles'].get(3, 0))
        bot_s3  = min(cibles_services, key=lambda sv: sv['cibles'].get(3, 0))
        _c = [
            f"{len(cibles_services)} service(s) ont des tâches planifiées dans ce PTA. "
            f"En moyenne, {avg_t1s:.1f} % des tâches devraient être exécutées à fin T1 "
            f"et {avg_t3s:.1f} % à fin T3."
        ]
        _c.append(
            f"Le service le plus avancé dès T1 est {top_s1['code']} "
            f"({top_s1['direction_code']}, {top_s1['cibles'].get(1, 0):.1f} %), "
            f"reflétant des activités concentrées en début d'année."
        )
        if len(cibles_services) > 1 and bot_s1['code'] != top_s1['code']:
            _c.append(
                f"Le service {bot_s1['code']} ({bot_s1['direction_code']}) "
                f"n'atteint que {bot_s1['cibles'].get(1, 0):.1f} % en T1, "
                f"avec des tâches planifiées plutôt en seconde partie d'année."
            )
        if len(cibles_services) > 1 and (top_s3['code'] != top_s1['code'] or bot_s3['code'] != bot_s1['code']):
            _c.append(
                f"À fin T3 : {top_s3['code']} est le plus avancé ({top_s3['cibles'].get(3, 0):.1f} %) "
                f"et {bot_s3['code']} le moins ({bot_s3['cibles'].get(3, 0):.1f} %)."
            )
        interp(doc, ' '.join(_c))

    return doc


# ── Export Excel statistiques ─────────────────────────────────────────────────

def _build_excel(annee, s):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    thin = Side(style='thin')
    brd  = Border(left=thin, right=thin, top=thin, bottom=thin)
    ctr  = Alignment(horizontal='center', vertical='center', wrap_text=True)
    lft  = Alignment(horizontal='left',   vertical='center', wrap_text=True)

    F_TITRE  = PatternFill("solid", fgColor="1F4E79")
    F_HDR    = PatternFill("solid", fgColor="2E75B6")
    F_INVEST = PatternFill("solid", fgColor="C0392B")
    F_FONCT  = PatternFill("solid", fgColor="27AE60")
    F_PROG   = PatternFill("solid", fgColor="9DC3E6")
    F_PROJ   = PatternFill("solid", fgColor="DEEAF1")
    F_TOTAL  = PatternFill("solid", fgColor="AED6F1")

    def _titre(ws, texte):
        ws.merge_cells('A1:J1')
        ws['A1'] = texte
        ws['A1'].font      = Font(bold=True, size=13, color="FFFFFF")
        ws['A1'].fill      = F_TITRE
        ws['A1'].alignment = ctr

    def _entete(ws, cols, row=2, fill=None):
        fill = fill or F_HDR
        for ci, col in enumerate(cols, 1):
            c = ws.cell(row=row, column=ci, value=col)
            c.font      = Font(bold=True, color="FFFFFF", size=9)
            c.fill      = fill
            c.alignment = ctr
            c.border    = brd
        return row + 1

    def _cell(ws, r, c, v, fill=None, bold=False, align=None):
        cell = ws.cell(row=r, column=c, value=v)
        if fill:  cell.fill = fill
        if bold:  cell.font = Font(bold=True, size=9)
        else:     cell.font = Font(size=9)
        cell.alignment = align or lft
        cell.border    = brd
        return cell

    def _num(v):
        try: return round(float(v or 0))
        except: return 0

    # ── Feuille 1 : Résumé général ─────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Résumé général"
    _titre(ws1, f"STATISTIQUES PTA {annee.annee} — Mairie d'Adja-Ouèrè — Résumé général")

    donnees_resume = [
        ("Budget total PTA",             f"{_fmt(s['total_budget'])} F CFA"),
        ("dont Investissement",           f"{_fmt(s['inv_budget'])} F CFA"),
        ("dont Fonctionnement",           f"{_fmt(s['fct_budget'])} F CFA"),
        ("",                              ""),
        ("Ressources propres (RP)",       f"{_fmt(s['total_rp'])} F CFA"),
        ("FADEC Affecté (FA)",            f"{_fmt(s['total_fa'])} F CFA"),
        ("FADEC Non Affecté (FNA)",       f"{_fmt(s['total_fn'])} F CFA"),
        ("PTFs / Partenaires (AP)",       f"{_fmt(s['total_ap'])} F CFA"),
        ("Autres Fonds (AF)",             f"{_fmt(s['total_af'])} F CFA"),
        ("",                              ""),
        ("Nombre de Programmes",          s['nb_programmes']),
        ("Nombre de Projets",             s['nb_projets']),
        ("Nombre d'Activités (total)",    s['nb_activites']),
        ("  Activités d'investissement",  s['nb_activites_invest']),
        ("  Activités de fonctionnement", s['nb_activites_fonct']),
        ("Nombre de Tâches (total)",      s['nb_taches']),
        ("  Tâches (investissement)",     s['nb_taches_invest']),
        ("  Tâches (fonctionnement)",     s['nb_taches_fonct']),
    ]
    r = 2
    for libelle, valeur in donnees_resume:
        if libelle == "":
            r += 1
            continue
        ws1.cell(row=r, column=1, value=libelle).font  = Font(size=9, bold=("dont" in libelle or "Activités d'" in libelle or "Tâches (" in libelle))
        ws1.cell(row=r, column=1).border    = brd
        ws1.cell(row=r, column=1).alignment = lft
        c2 = ws1.cell(row=r, column=2, value=valeur)
        c2.font      = Font(size=9, bold=True)
        c2.border    = brd
        c2.alignment = ctr
        r += 1

    ws1.column_dimensions['A'].width = 38
    ws1.column_dimensions['B'].width = 26

    # ── Feuille 2 : Par Programme ──────────────────────────────────────────
    ws2 = wb.create_sheet("Par Programme")
    _titre(ws2, f"STATISTIQUES PAR PROGRAMME — PTA {annee.annee}")
    cols2 = ["Code", "Programme / Projet", "Nb Projets", "Nb Activités", "Nb Act. Invest.",
             "Nb Act. Fonct.", "Nb Tâches", "Budget Total (F CFA)",
             "Ressources propres", "FADEC Affecté", "FADEC Non Aff.", "PTFs", "Autres Fonds"]
    next_r = _entete(ws2, cols2, row=2)

    total_p_acts = total_p_budget = 0
    for pe in s['stats_par_programme']:
        _cell(ws2, next_r, 1, pe['code'],         F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r, 2, pe['nom'],           F_PROG, bold=True)
        _cell(ws2, next_r, 3, pe['nb_projets'],    F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r, 4, pe['nb_activites'],  F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r, 5, pe['nb_activites_invest'], F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r, 6, pe['nb_activites_fonct'],  F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r, 7, pe['nb_taches'],     F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r, 8, _num(pe['budget']),  F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r, 9, _num(pe['rp']),      F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r,10, _num(pe['fa']),      F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r,11, _num(pe['fn']),      F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r,12, _num(pe['ap']),      F_PROG, bold=True, align=ctr)
        _cell(ws2, next_r,13, _num(pe['af']),      F_PROG, bold=True, align=ctr)
        total_p_acts   += pe['nb_activites']
        total_p_budget += _num(pe['budget'])
        next_r += 1
        for prj in pe.get('projets', []):
            _cell(ws2, next_r, 1, prj['code'],              None, align=ctr)
            _cell(ws2, next_r, 2, f"  {prj['nom']}",       F_PROJ)
            _cell(ws2, next_r, 3, "",                       F_PROJ, align=ctr)
            _cell(ws2, next_r, 4, prj['nb_activites'],      F_PROJ, align=ctr)
            _cell(ws2, next_r, 5, prj['nb_activites_invest'],F_PROJ, align=ctr)
            _cell(ws2, next_r, 6, prj['nb_activites_fonct'], F_PROJ, align=ctr)
            _cell(ws2, next_r, 7, prj['nb_taches'],         F_PROJ, align=ctr)
            _cell(ws2, next_r, 8, _num(prj['budget']),      F_PROJ, align=ctr)
            _cell(ws2, next_r, 9, _num(prj['rp']),          F_PROJ, align=ctr)
            _cell(ws2, next_r,10, _num(prj['fa']),          F_PROJ, align=ctr)
            _cell(ws2, next_r,11, _num(prj['fn']),          F_PROJ, align=ctr)
            _cell(ws2, next_r,12, _num(prj['ap']),          F_PROJ, align=ctr)
            _cell(ws2, next_r,13, _num(prj['af']),          F_PROJ, align=ctr)
            next_r += 1

    _cell(ws2, next_r, 1, "TOTAL",             F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r, 2, "",                  F_TOTAL)
    _cell(ws2, next_r, 3, "",                  F_TOTAL, align=ctr)
    _cell(ws2, next_r, 4, total_p_acts,        F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r, 5, s['nb_activites_invest'], F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r, 6, s['nb_activites_fonct'],  F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r, 7, s['nb_taches'],      F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r, 8, total_p_budget,      F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r, 9, _num(s['total_rp']), F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r,10, _num(s['total_fa']), F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r,11, _num(s['total_fn']), F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r,12, _num(s['total_ap']), F_TOTAL, bold=True, align=ctr)
    _cell(ws2, next_r,13, _num(s['total_af']), F_TOTAL, bold=True, align=ctr)

    widths2 = [10, 45, 10, 12, 14, 14, 10, 20, 18, 16, 16, 14, 14]
    for i, w in enumerate(widths2, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.freeze_panes = 'A3'

    # ── Feuille 3 : Par Direction ──────────────────────────────────────────
    ws3 = wb.create_sheet("Par Direction")
    _titre(ws3, f"STATISTIQUES PAR DIRECTION RESPONSABLE — PTA {annee.annee}")
    cols3 = ["Code", "Direction", "Nb Act. (Total)", "Nb Act. Invest.", "Nb Act. Fonct.",
             "Nb Tâches (Total)"]
    next_r = _entete(ws3, cols3, row=2)

    for d in s['stats_directions']:
        d_inv = next((x for x in s['stats_directions_invest'] if x['code'] == d['code']), {})
        d_fct = next((x for x in s['stats_directions_fonct']  if x['code'] == d['code']), {})
        _cell(ws3, next_r, 1, d['code'],                  None, align=ctr)
        _cell(ws3, next_r, 2, d['nom'])
        _cell(ws3, next_r, 3, d['nb_activites'],          None, align=ctr)
        _cell(ws3, next_r, 4, d_inv.get('nb_activites',0),None, align=ctr)
        _cell(ws3, next_r, 5, d_fct.get('nb_activites',0),None, align=ctr)
        _cell(ws3, next_r, 6, d['nb_taches'],             None, align=ctr)
        next_r += 1

    for i, w in enumerate([10, 45, 16, 16, 16, 16], 1):
        ws3.column_dimensions[get_column_letter(i)].width = w
    ws3.freeze_panes = 'A3'

    # ── Feuille 4 : Activités d'investissement ─────────────────────────────
    ws4 = wb.create_sheet("Activités investissement")
    _titre(ws4, f"LISTE DES ACTIVITÉS D'INVESTISSEMENT — PTA {annee.annee}")
    cols4 = ["Code", "Désignation", "Direction", "Période", "Poids (%)",
             "Budget (F CFA)", "Ressources propres", "FADEC Affecté",
             "FADEC Non Aff.", "PTFs", "Autres Fonds", "Nb Tâches"]
    next_r = _entete(ws4, cols4, row=2, fill=F_INVEST)

    for ai in s['activites_investissement']:
        _cell(ws4, next_r,  1, ai['code'],             None, align=ctr)
        _cell(ws4, next_r,  2, ai['nom'])
        _cell(ws4, next_r,  3, ai['direction'],        None, align=ctr)
        _cell(ws4, next_r,  4, ai['periode'],          None, align=ctr)
        _cell(ws4, next_r,  5, round(ai['poids'],2),   None, align=ctr)
        _cell(ws4, next_r,  6, _num(ai['budget']),     None, align=ctr)
        _cell(ws4, next_r,  7, _num(ai['rp']),         None, align=ctr)
        _cell(ws4, next_r,  8, _num(ai['fa']),         None, align=ctr)
        _cell(ws4, next_r,  9, _num(ai['fn']),         None, align=ctr)
        _cell(ws4, next_r, 10, _num(ai['ap']),         None, align=ctr)
        _cell(ws4, next_r, 11, _num(ai['af']),         None, align=ctr)
        _cell(ws4, next_r, 12, ai['nb_taches'],        None, align=ctr)
        next_r += 1

    tot_inv_budget = sum(_num(ai['budget']) for ai in s['activites_investissement'])
    _cell(ws4, next_r,  1, "TOTAL",          F_TOTAL, bold=True, align=ctr)
    _cell(ws4, next_r,  2, "",               F_TOTAL)
    _cell(ws4, next_r,  3, "",               F_TOTAL, align=ctr)
    _cell(ws4, next_r,  4, "",               F_TOTAL, align=ctr)
    _cell(ws4, next_r,  5, "",               F_TOTAL, align=ctr)
    _cell(ws4, next_r,  6, tot_inv_budget,   F_TOTAL, bold=True, align=ctr)
    _cell(ws4, next_r,  7, _num(s['inv_rp']),F_TOTAL, bold=True, align=ctr)
    _cell(ws4, next_r,  8, _num(s['inv_fa']),F_TOTAL, bold=True, align=ctr)
    _cell(ws4, next_r,  9, _num(s['inv_fn']),F_TOTAL, bold=True, align=ctr)
    _cell(ws4, next_r, 10, _num(s['inv_ap']),F_TOTAL, bold=True, align=ctr)
    _cell(ws4, next_r, 11, _num(s['inv_af']),F_TOTAL, bold=True, align=ctr)
    _cell(ws4, next_r, 12, s['nb_taches_invest'], F_TOTAL, bold=True, align=ctr)

    for i, w in enumerate([12, 50, 12, 18, 10, 18, 18, 16, 16, 14, 14, 10], 1):
        ws4.column_dimensions[get_column_letter(i)].width = w
    ws4.freeze_panes = 'A3'

    return wb


# ── Routes ────────────────────────────────────────────────────────────────────

@stats_bp.route('/')
@login_required
def index():
    annee = get_annee()
    if not annee:
        return render_template('stats/index.html', annee=None)
    s = _compute_stats(annee)
    cibles = _compute_cibles(s['programmes'])
    cibles_directions = _compute_cibles_directions(annee, s['directions'])
    all_services = Service.query.order_by(Service.nom).all()
    cibles_services = _compute_cibles_services(annee, all_services)
    return render_template('stats/index.html', annee=annee, cibles=cibles,
                           cibles_directions=cibles_directions,
                           cibles_services=cibles_services, **s)


@stats_bp.route('/rapport/word')
@login_required
def rapport_word():
    annee = get_annee()
    if not annee:
        flash('Aucune année active.', 'danger')
        return redirect(url_for('stats.index'))
    try:
        s                 = _compute_stats(annee)
        cibles            = _compute_cibles(s['programmes'])
        cibles_directions = _compute_cibles_directions(annee, s['directions'])
        all_services      = Service.query.order_by(Service.nom).all()
        cibles_services   = _compute_cibles_services(annee, all_services)
        static_img        = os.path.join(current_app.root_path, 'static', 'img')
        doc = _build_word(annee, s, static_img,
                          cibles=cibles,
                          cibles_directions=cibles_directions,
                          cibles_services=cibles_services)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        fname = f"Rapport_Stats_PTA_{annee.annee}.docx"
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=fname,
        )
    except Exception as e:
        flash(f'Erreur lors de la génération du rapport : {e}', 'danger')
        return redirect(url_for('stats.index'))


@stats_bp.route('/rapport/excel')
@login_required
def rapport_excel():
    annee = get_annee()
    if not annee:
        flash('Aucune année active.', 'danger')
        return redirect(url_for('stats.index'))
    try:
        s = _compute_stats(annee)
        wb = _build_excel(annee, s)
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        fname = f"Statistiques_PTA_{annee.annee}.xlsx"
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=fname,
        )
    except Exception as e:
        flash(f'Erreur lors de la génération Excel : {e}', 'danger')
        return redirect(url_for('stats.index'))
